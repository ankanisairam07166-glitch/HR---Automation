from flask import Flask, request, jsonify, redirect, render_template_string, send_from_directory
from flask_cors import CORS
from datetime import datetime, timedelta, timezone
from db import Candidate, SessionLocal
import threading
import asyncio
import time
import traceback
import os
import json
from sqlalchemy import func, and_
import requests
import logging
from logging.handlers import RotatingFileHandler
from functools import wraps
from tenacity import retry, stop_after_attempt, wait_exponential
import sys
from flask_caching import Cache
import redis
from concurrent.futures import ThreadPoolExecutor
import uuid
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from interview_automation import start_interview_automation, stop_interview_automation
from werkzeug.utils import secure_filename

# Import your existing modules
try:
    from scraper import scrape_job
    from latest import create_programming_assessment
    from test_link import get_invite_link
    from clint_recruitment_system import run_recruitment_with_invite_link
    from email_util import send_assessment_email, send_assessment_reminder, send_interview_confirmation_email, send_interview_link_email, send_rejection_email
except ImportError as e:
    logging.error(f"Critical module import failed: {e}")
    raise

# Add after existing imports
try:
    from testlify_results_scraper import scrape_all_pending_assessments, scrape_assessment_results_by_name
except ImportError as e:
    logging.warning(f"Testlify scraper not available: {e}")

def _ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path

def _append_jsonl(path: str, obj: dict) -> None:
    _ensure_dir(os.path.dirname(path))
    with open(path, "a", encoding="utf-8") as fh:
        fh.write(json.dumps(obj, ensure_ascii=False) + "\n")

def _ok_preflight():
    # If you have CORS set globally, 200 is fine for OPTIONS
    return "", 200

ALLOWED_EXTS = {"webm", "mp4", "mkv", "mov"}

def _ext_from_filename(name: str, default_ext: str = "webm") -> str:
    ext = os.path.splitext(name or "")[1].lower().lstrip(".")
    return ext if ext in ALLOWED_EXTS else default_ext

# Setup proper logging
def setup_logging():
    """Configure logging for production"""
    if not os.path.exists('logs'):
        os.makedirs('logs')
    
    # Create formatter
    formatter = logging.Formatter(
        '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
    )
    
    # File handler
    file_handler = RotatingFileHandler(
        'logs/talentflow.log',
        maxBytes=10485760,  # 10MB
        backupCount=10,
        encoding='utf-8'
    )
    file_handler.setFormatter(formatter)
    file_handler.setLevel(logging.INFO)
    
    # Console handler  
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.DEBUG)
    console_handler.setFormatter(formatter)
    
    # Configure root logger
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    # Clear existing handlers
    for handler in logger.handlers[:]:
        logger.removeHandler(handler)
    
    # Add our handlers
    logger.addHandler(file_handler)
    
    # Only add console handler in development
    if os.getenv('FLASK_ENV') == 'development':
        logger.addHandler(console_handler)
    
    return logger

logger = setup_logging()

# Configuration from environment
ASSESSMENT_CONFIG = {
    'EXPIRY_HOURS': int(os.getenv('ASSESSMENT_EXPIRY_HOURS', '48')),
    'REMINDER_HOURS': int(os.getenv('ASSESSMENT_REMINDER_HOURS', '24')),
    'INTERVIEW_DELAY_DAYS': int(os.getenv('INTERVIEW_DELAY_DAYS', '3')),
    'ATS_THRESHOLD': float(os.getenv('ATS_THRESHOLD', '70')),
    'MAX_RETRIES': int(os.getenv('MAX_RETRIES', '3')),
    'RETRY_DELAY': int(os.getenv('RETRY_DELAY', '2'))
}

# Create Flask app
app = Flask(__name__)

# Setup caching
cache_config = {
    'CACHE_TYPE': 'simple',  # Use Redis in production
    'CACHE_DEFAULT_TIMEOUT': 300  # 5 minutes
}

if os.getenv('REDIS_URL'):
    cache_config = {
        'CACHE_TYPE': 'redis',
        'CACHE_REDIS_URL': os.getenv('REDIS_URL'),
        'CACHE_DEFAULT_TIMEOUT': 300
    }

cache = Cache(app, config=cache_config)

# Enhanced CORS Configuration
CORS(app, 
     origins=["http://localhost:3000", "http://127.0.0.1:3000", "https://yourfrontenddomain.com","http://127.0.0.1:3001"],
     allow_headers=["Content-Type", "Authorization", "X-Requested-With", "Accept", "Cache-Control","X-Api-Key"],
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
     supports_credentials=True,
     expose_headers=["Content-Type", "Authorization"]),

# NEXTJS_URL = os.getenv('NEXTJS_URL', 'http://localhost:3001')
HEYGEN_API_URL = "https://api.heygen.com/v1/your-avatar-endpoint"  # Change this!

HEYGEN_API_KEY = os.getenv('HEYGEN_API_KEY', '')
# Thread pool for background tasks
executor = ThreadPoolExecutor(max_workers=4)

# Pipeline status tracking
pipeline_status = {}
pipeline_lock = threading.Lock()

# Performance monitoring
request_metrics = {
    'total_requests': 0,
    'avg_response_time': 0,
    'slow_requests': 0
}


#  Admin notification function
def notify_admin(subject, message, error_details=None):
    """Send critical notifications to admin"""
    try:
        admin_email = os.getenv('ADMIN_EMAIL')
        if not admin_email:
            logger.warning("ADMIN_EMAIL not set, skipping notification")
            return
        
        from email_util import send_email
        
        body_html = f"""
        <html>
            <body>
                <h2>TalentFlow AI Alert: {subject}</h2>
                <p>{message}</p>
                {f'<pre>{error_details}</pre>' if error_details else ''}
                <p>Time: {datetime.now().isoformat()}</p>
            </body>
        </html>
        """
        
        send_email(admin_email, f"[TalentFlow Alert] {subject}", body_html)
        
    except Exception as e:
        logger.error(f"Failed to send admin notification: {e}")

# Add request timing middleware
@app.before_request
def before_request():
    request.start_time = time.time()
    request_metrics['total_requests'] += 1

@app.after_request
def after_request(response):
    if hasattr(request, 'start_time'):
        duration = time.time() - request.start_time
        
        # Update metrics
        if request_metrics['avg_response_time'] == 0:
            request_metrics['avg_response_time'] = duration
        else:
            request_metrics['avg_response_time'] = (request_metrics['avg_response_time'] + duration) / 2
        
        if duration > 5.0:  # Slow request threshold
            request_metrics['slow_requests'] += 1
            logger.warning(f"Slow request: {request.method} {request.path} took {duration:.2f}s")
        
        # Add performance headers
        response.headers['X-Response-Time'] = f"{duration:.3f}s"
        response.headers['X-Request-ID'] = getattr(request, 'request_id', 'unknown')
    
    return response

# Add request logging for debugging
@app.before_request
def log_request_info():
    """Log incoming requests for debugging"""
    request.request_id = str(uuid.uuid4())[:8]
    if request.endpoint and (request.endpoint.startswith('api_') or 'api' in request.path):
        logger.info(f"🌐 [{request.request_id}] {request.method} {request.path} from {request.remote_addr}")
        if request.method == 'OPTIONS':
            logger.info(f"🔧 [{request.request_id}] CORS preflight for {request.path}")

# Rate limiting decorator with better performance
def rate_limit(max_calls=10, time_window=60):
    """Enhanced rate limiting decorator with memory optimization"""
    calls = {}
    cleanup_counter = 0
    
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            nonlocal cleanup_counter
            
            # Skip rate limiting for OPTIONS requests (CORS preflight)
            if request.method == 'OPTIONS':
                return func(*args, **kwargs)
            
            now = time.time()
            key = request.remote_addr
            
            # Periodic cleanup to prevent memory leaks
            cleanup_counter += 1
            if cleanup_counter % 100 == 0:
                cutoff = now - time_window * 2
                for ip in list(calls.keys()):
                    calls[ip] = [call_time for call_time in calls.get(ip, []) if call_time > cutoff]
                    if not calls[ip]:
                        del calls[ip]
            
            if key not in calls:
                calls[key] = []
            
            # Remove old calls
            calls[key] = [call_time for call_time in calls[key] if now - call_time < time_window]
            
            if len(calls[key]) >= max_calls:
                logger.warning(f"Rate limit exceeded for {key}")
                return jsonify({"error": "Rate limit exceeded"}), 429
            
            calls[key].append(now)
            return func(*args, **kwargs)
        return wrapper
    return decorator

# Pipeline management functions
def update_pipeline_status(job_id, status, message, progress=None):
    """Thread-safe pipeline status updates"""
    with pipeline_lock:
        pipeline_status[str(job_id)] = {
            'status': status,
            'message': message,
            'progress': progress,
            'timestamp': datetime.now().isoformat(),
            'job_id': str(job_id)
        }
    logger.info(f"Pipeline {job_id}: {status} - {message}")

def get_pipeline_status(job_id=None):
    """Get pipeline status (thread-safe)"""
    with pipeline_lock:
        if job_id:
            return pipeline_status.get(str(job_id))
        return dict(pipeline_status)

# Optimized data fetching with caching
@cache.memoize(timeout=300)
def get_cached_jobs():
    """Cached job fetching"""
    try:
        API_KEY = os.getenv("BAMBOOHR_API_KEY")
        SUBDOMAIN = os.getenv("BAMBOOHR_SUBDOMAIN")
        
        if not API_KEY or not SUBDOMAIN:
            raise ValueError("BambooHR credentials not configured")
            
        auth = (API_KEY, "x")
        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        url = f"https://api.bamboohr.com/api/gateway.php/{SUBDOMAIN}/v1/applicant_tracking/jobs/"
        
        resp = requests.get(url, auth=auth, headers=headers, timeout=10)
        resp.raise_for_status()
        
        jobs = resp.json()
        open_jobs = []
        
        session = SessionLocal()
        try:
            for job in jobs:
                if job.get("status", {}).get("label", "").lower() == "open":
                    # Get candidate count for this job
                    candidate_count = session.query(Candidate).filter_by(job_id=str(job["id"])).count()
                    
                    open_jobs.append({
                        "id": job["id"],
                        "title": job.get("title", {}).get("label", ""),
                        "location": job.get("location", {}).get("label", ""),
                        "department": job.get("department", {}).get("label", ""),
                        "postingUrl": job.get("postingUrl", ""),
                        "applications": candidate_count,
                        "status": "Active",
                        "description": job.get("description", "")
                    })
        finally:
            session.close()
        
        return open_jobs
        
    except Exception as e:
        logger.error(f"BambooHR API error: {e}")
        # Fallback to database
        return get_jobs_from_database()

def get_jobs_from_database():
    """Fallback job fetching from database"""
    session = SessionLocal()
    try:
        jobs_data = session.query(
            Candidate.job_id,
            Candidate.job_title,
            func.count(Candidate.id).label('applications')
        ).filter(
            Candidate.job_id.isnot(None),
            Candidate.job_title.isnot(None)
        ).group_by(
            Candidate.job_id, 
            Candidate.job_title
        ).all()
        
        jobs = []
        for job_id, job_title, app_count in jobs_data:
            jobs.append({
                'id': str(job_id),
                'title': job_title,
                'department': 'Engineering',
                'location': 'Remote',
                'applications': app_count,
                'status': 'Active',
                'description': f'Job description for {job_title}',
                'postingUrl': ''
            })
        
        return jobs
    finally:
        session.close()

@cache.memoize(timeout=180)  # 3 minutes cache
def get_cached_candidates(job_id=None, status_filter=None):
    """Cached candidate fetching with optimized queries"""
    session = SessionLocal()
    try:
        query = session.query(Candidate)
        
        if job_id:
            query = query.filter_by(job_id=str(job_id))
        
        if status_filter:
            query = query.filter_by(status=status_filter)
        
        # Optimize query - only get needed fields for list view
        candidates = query.all()
        
        result = []
        for c in candidates:
            try:
                # Calculate time remaining for assessment
                time_remaining = None
                link_expired = False
                
                if c.exam_link_sent_date and not c.exam_completed:
                    deadline = c.exam_link_sent_date + timedelta(hours=ASSESSMENT_CONFIG['EXPIRY_HOURS'])
                    if datetime.now() < deadline:
                        time_remaining = (deadline - datetime.now()).total_seconds() / 3600
                    else:
                        link_expired = True
                
                candidate_data = {
                    "id": c.id,
                    "name": c.name or "Unknown",
                    "email": c.email or "",
                    "job_id": c.job_id,
                    "job_title": c.job_title or "Unknown Position",
                    "status": c.status,
                    "ats_score": float(c.ats_score) if c.ats_score else 0.0,
                    "linkedin": c.linkedin,
                    "github": c.github,
                    "phone": getattr(c, 'phone', None),
                    "resume_path": c.resume_path,
                    "processed_date": c.processed_date.isoformat() if c.processed_date else None,
                    "score_reasoning": c.score_reasoning,
                    
                    # Assessment fields
                    "assessment_invite_link": c.assessment_invite_link,
                    "exam_link_sent": bool(c.exam_link_sent),
                    "exam_link_sent_date": c.exam_link_sent_date.isoformat() if c.exam_link_sent_date else None,
                    "exam_completed": bool(c.exam_completed),
                    "exam_completed_date": c.exam_completed_date.isoformat() if c.exam_completed_date else None,
                    "link_expired": link_expired,
                    "time_remaining_hours": time_remaining,
                    
                    # Exam results
                    "exam_percentage": float(c.exam_percentage) if c.exam_percentage else None,
                    
                    # Interview fields
                    "interview_scheduled": bool(c.interview_scheduled),
                    "interview_date": c.interview_date.isoformat() if c.interview_date else None,
                    "interview_link": c.interview_link,
                    
                    # Status fields
                    "final_status": c.final_status,
                }
                
                result.append(candidate_data)
                
            except Exception as e:
                logger.error(f"Error processing candidate {c.id}: {e}")
                continue
        
        return result
    finally:
        session.close()

@app.route('/', methods=['GET'])
def home():
    """Enhanced root endpoint with comprehensive API information"""
    try:
        # Get system statistics
        session = SessionLocal()
        try:
            stats = {
                "total_candidates": session.query(Candidate).count(),
                "total_jobs": session.query(Candidate.job_id).distinct().count(),
                "shortlisted_candidates": session.query(Candidate).filter_by(status='Shortlisted').count(),
                "completed_assessments": session.query(Candidate).filter_by(exam_completed=True).count(),
                "scheduled_interviews": session.query(Candidate).filter_by(interview_scheduled=True).count(),
            }
        except Exception:
            stats = {"error": "Could not fetch statistics"}
        finally:
            session.close()
        
        # Check system health
        system_health = {
            "database": "healthy",
            "cache": "healthy",
            "interview_automation": "running" if hasattr(app, 'interview_automation') else "stopped"
        }
        
        # API documentation with examples
        api_docs = {
            "endpoints": {
                "GET /api/jobs": {
                    "description": "Get all job postings",
                    "example": f"{request.host_url}api/jobs"
                },
                "GET /api/candidates": {
                    "description": "Get candidates (filterable by job_id, status)",
                    "example": f"{request.host_url}api/candidates?job_id=123&status=Shortlisted"
                },
                "POST /api/run_full_pipeline": {
                    "description": "Start recruitment pipeline for a job",
                    "example_payload": {
                        "job_id": "123",
                        "job_title": "Software Engineer",
                        "job_desc": "Job description here"
                    }
                },
                "GET /api/pipeline_status/<job_id>": {
                    "description": "Check pipeline status for specific job",
                    "example": f"{request.host_url}api/pipeline_status/123"
                }
            }
        }
        
        return jsonify({
            "message": "🚀 TalentFlow AI Backend API",
            "tagline": "Intelligent Recruitment Automation Platform",
            "version": "2.1.0",
            "status": "operational",
            "timestamp": datetime.now().isoformat(),
            "uptime": "System started successfully",
            
            # System Statistics
            "statistics": stats,
            "system_health": system_health,
            
            # Quick Links
            "quick_links": {
                "health_check": f"{request.host_url}health",
                "api_documentation": f"{request.host_url}api/routes",
                "frontend_dashboard": "http://localhost:3000"
            },
            
            # API Information
            "api_info": api_docs,
            
            # Contact & Support
            "support": {
                "company": os.getenv('COMPANY_NAME', 'TalentFlow AI'),
                "admin_email": os.getenv('ADMIN_EMAIL', 'admin@talentflow.ai'),
                "documentation": "https://docs.talentflow.ai"
            },
            
            # Features Highlight
            "features": [
                "🤖 AI-Powered Resume Screening",
                "📝 Automated Assessment Creation", 
                "📧 Smart Email Automation",
                "📊 Real-time Analytics Dashboard",
                "🎥 AI Avatar Interviews",
                "⚡ Pipeline Automation"
            ]
        }), 200
        
    except Exception as e:
        logger.error(f"Error in enhanced home route: {e}")
        return jsonify({
            "message": "🚀 TalentFlow AI Backend API",
            "version": "2.1.0",
            "status": "running",
            "error": "Partial system information available",
            "basic_endpoints": ["/api/jobs", "/api/candidates", "/health"]
        }), 200

# Enhanced API endpoints
@app.route('/api/jobs', methods=['GET', 'OPTIONS'])
@rate_limit(max_calls=30, time_window=60)
def api_jobs():
    """Enhanced API endpoint to get jobs with caching"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        jobs = get_cached_jobs()
        return jsonify(jobs), 200
    except Exception as e:
        logger.error(f"Error in api_jobs: {e}", exc_info=True)
        return jsonify({"error": "Failed to fetch jobs", "message": str(e)}), 500

@app.route('/api/candidates', methods=['GET', 'OPTIONS'])
@rate_limit(max_calls=60, time_window=60)
def api_candidates():
    """Enhanced API endpoint to get candidates with caching"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        job_id = request.args.get('job_id')
        status_filter = request.args.get('status')
        
        candidates = get_cached_candidates(job_id, status_filter)
        return jsonify(candidates), 200
        
    except Exception as e:
        logger.error(f"Error in api_candidates: {e}", exc_info=True)
        return jsonify({"error": "Failed to fetch candidates", "message": str(e)}), 500

@app.route('/api/run_full_pipeline', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=5, time_window=300)
def api_run_full_pipeline():
    """Enhanced pipeline API with status tracking"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        job_id = data.get('job_id')
        job_title = data.get('job_title')
        job_desc = data.get('job_desc', "")
        
        logger.info(f"[{request.request_id}] Pipeline request: job_id={job_id}, job_title={job_title}")
        
        if not job_id or not job_title:
            return jsonify({"success": False, "message": "job_id and job_title are required"}), 400
        
        # Check if pipeline is already running for this job
        current_status = get_pipeline_status(job_id)
        if current_status and current_status.get('status') == 'running':
            return jsonify({
                "success": False,
                "message": f"Pipeline already running for {job_title}",
                "status": current_status
            }), 409
        
        # # Update status to starting
        # update_pipeline_status(job_id, 'starting', f'Initializing pipeline for {job_title}', 0)
        
# Update status to starting
        update_pipeline_status(job_id, 'starting', f'Initializing pipeline for {job_title}', 0)
        
        # Start the pipeline in background thread
        future = executor.submit(run_pipeline_with_monitoring, job_id, job_title, job_desc)
        
        # Store future for tracking
        with pipeline_lock:
            pipeline_status[str(job_id)]['future'] = future
        
        return jsonify({
            "success": True, 
            "message": f"Pipeline started for {job_title}",
            "job_id": job_id,
            "estimated_time": "5-10 minutes",
            "status_endpoint": f"/api/pipeline_status/{job_id}"
        }), 200
        
    except Exception as e:
        logger.error(f"Error in run_full_pipeline: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/pipeline_status', methods=['GET', 'OPTIONS'])
@app.route('/api/pipeline_status/<job_id>', methods=['GET', 'OPTIONS'])
def api_pipeline_status(job_id=None):
    """Get pipeline status for specific job or all jobs"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        if job_id:
            status = get_pipeline_status(job_id)
            if not status:
                return jsonify({"success": False, "message": "Pipeline not found"}), 404
            
            # Clean up the status (remove future object for JSON serialization)
            clean_status = {k: v for k, v in status.items() if k != 'future'}
            return jsonify({"success": True, "status": clean_status}), 200
        else:
            all_status = get_pipeline_status()
            # Clean up all statuses
            clean_statuses = {k: {sk: sv for sk, sv in v.items() if sk != 'future'} 
                            for k, v in all_status.items()}
            return jsonify({"success": True, "pipelines": clean_statuses}), 200
            
    except Exception as e:
        logger.error(f"Error in pipeline_status: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

def run_pipeline_with_monitoring(job_id, job_title, job_desc):
    """Enhanced pipeline runner with detailed progress tracking"""
    start_time = time.time()
    
    try:
        logger.info(f"Starting monitored pipeline for job_id={job_id}")
        update_pipeline_status(job_id, 'running', 'Pipeline started', 10)
        
        # Clear relevant caches
        cache.delete_memoized(get_cached_candidates)
        cache.delete_memoized(get_cached_jobs)
        
        full_recruitment_pipeline(job_id, job_title, job_desc)
        
        duration = time.time() - start_time
        update_pipeline_status(job_id, 'completed', f'Pipeline completed successfully in {duration:.1f}s', 100)
        logger.info(f"Pipeline completed successfully in {duration:.2f} seconds")
        
    except Exception as e:
        duration = time.time() - start_time
        error_msg = f"Pipeline failed after {duration:.2f} seconds: {str(e)}"
        update_pipeline_status(job_id, 'error', error_msg, None)
        logger.error(error_msg, exc_info=True)

def full_recruitment_pipeline(job_id, job_title, job_desc):
    """Enhanced recruitment pipeline with progress tracking"""
    try:
        logger.info(f"Starting full recruitment pipeline for job_id={job_id}, job_title={job_title}")
        
        # STEP 1: Scraping (20% progress)
        try:
            update_pipeline_status(job_id, 'running', 'Scraping resumes...', 20)
            logger.info(f"STEP 1: Scraping resumes for job_id={job_id}")
            asyncio.run(scrape_job(job_id))
            logger.info("✅ Scraping completed successfully")
        except Exception as e:
            logger.error(f"Scraping failed: {str(e)}", exc_info=True)
            # Continue with next step even if scraping fails
        
        # STEP 2: Create assessment (40% progress)
        try:
            update_pipeline_status(job_id, 'running', 'Creating programming assessment...', 40)
            logger.info(f"STEP 2: Creating assessment for '{job_title}' in Testlify")
            create_programming_assessment(job_title, job_desc)
            logger.info("✅ Assessment created successfully")
        except Exception as e:
            logger.error(f"Assessment creation failed: {str(e)}", exc_info=True)
        
        # STEP 3: Get invite link (60% progress)
        invite_link = None
        try:
            update_pipeline_status(job_id, 'running', 'Extracting assessment invite link...', 60)
            logger.info(f"STEP 3: Extracting invite link for '{job_title}' from Testlify")
            invite_link = get_invite_link(job_title)
            if invite_link:
                logger.info(f"✅ Got invite link: {invite_link}")
        except Exception as e:
            logger.error(f"Invite link extraction failed: {str(e)}", exc_info=True)
        
        if not invite_link:
            invite_link = f"https://candidate.testlify.com/assessment/{job_id}"
            logger.warning(f"Using fallback invite link: {invite_link}")
        
        # STEP 4: Run AI screening (80% progress)
        try:
            update_pipeline_status(job_id, 'running', 'Running AI-powered screening...', 80)
            logger.info("STEP 4: Running AI-powered screening...")
            run_recruitment_with_invite_link(
                job_id=job_id, 
                job_title=job_title, 
                job_desc=job_desc, 
                invite_link=invite_link
            )
            logger.info("✅ AI screening completed successfully")
        except Exception as e:
            logger.error(f"AI screening failed: {str(e)}", exc_info=True)
            raise  # This is critical, so we raise
        
        # Final step: Clear caches
        cache.delete_memoized(get_cached_candidates)
        cache.delete_memoized(get_cached_jobs)
        
        logger.info("🚀 Recruitment pipeline finished successfully")
            
    except Exception as e:
        logger.error(f"Fatal pipeline error: {e}", exc_info=True)
        raise

@app.route('/api/recruitment-stats', methods=['GET', 'OPTIONS'])
@rate_limit(max_calls=20, time_window=60)
@cache.memoize(timeout=600)  # 10 minute cache
def api_recruitment_stats():
    """Cached recruitment statistics"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        stats = []
        current_date = datetime.now()
        
        # Get last 6 months of data efficiently
        for i in range(6):
            try:
                month_date = current_date - timedelta(days=30*i)
                month_name = month_date.strftime('%b')
                
                # Calculate month boundaries
                month_start = month_date.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                if month_start.month == 12:
                    month_end = month_start.replace(year=month_start.year + 1, month=1, day=1) - timedelta(seconds=1)
                else:
                    month_end = month_start.replace(month=month_start.month + 1, day=1) - timedelta(seconds=1)
                
                # Single query for all stats
                applications = session.query(func.count(Candidate.id)).filter(
                    and_(
                        Candidate.processed_date >= month_start,
                        Candidate.processed_date <= month_end
                    )
                ).scalar() or 0
                
                interviews = session.query(func.count(Candidate.id)).filter(
                    and_(
                        Candidate.interview_scheduled == True,
                        Candidate.interview_date >= month_start,
                        Candidate.interview_date <= month_end
                    )
                ).scalar() or 0
                
                hires = session.query(func.count(Candidate.id)).filter(
                    and_(
                        Candidate.final_status == "Hired",
                        Candidate.processed_date >= month_start,
                        Candidate.processed_date <= month_end
                    )
                ).scalar() or 0
                
                stats.append({
                    "month": month_name,
                    "applications": applications,
                    "interviews": interviews,
                    "hires": hires
                })
                
            except Exception as e:
                logger.error(f"Error calculating stats for month {i}: {e}")
                stats.append({
                    "month": (current_date - timedelta(days=30*i)).strftime('%b'),
                    "applications": 0,
                    "interviews": 0,
                    "hires": 0
                })
        
        # Reverse to get chronological order
        stats.reverse()
        
        logger.info(f"Generated recruitment stats for {len(stats)} months")
        return jsonify(stats), 200
        
    except Exception as e:
        logger.error(f"Error in api_recruitment_stats: {e}", exc_info=True)
        return jsonify({"error": "Failed to get statistics", "message": str(e)}), 500
    finally:
        session.close()

@app.route('/api/send_reminder/<int:candidate_id>', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=10, time_window=60)
def api_send_reminder(candidate_id):
    """Send reminder to specific candidate"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return jsonify({"success": False, "message": "Candidate not found"}), 404
            
            # Check if candidate is eligible for reminder
            if not candidate.exam_link_sent or candidate.exam_completed:
                return jsonify({"success": False, "message": "Candidate not eligible for reminder"}), 400
            
            # Calculate hours remaining
            hours_remaining = 24  # Default
            if candidate.exam_link_sent_date:
                deadline = candidate.exam_link_sent_date + timedelta(hours=48)
                hours_remaining = max(0, int((deadline - datetime.now()).total_seconds() / 3600))
            
            # Send reminder email
            send_assessment_reminder(candidate, hours_remaining)
            
            # Update reminder tracking
            candidate.reminder_sent = True
            candidate.reminder_sent_date = datetime.now()
            session.commit()
            
            # Clear cache
            cache.delete_memoized(get_cached_candidates)
            
            return jsonify({
                "success": True,
                "message": f"Reminder sent to {candidate.name}"
            }), 200
            
        finally:
            session.close()
        
    except Exception as e:
        logger.error(f"Error in send_reminder: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/assessments', methods=['GET'])
def api_assessments():
    return jsonify([]), 200


# COPY THIS EXACTLY AND PASTE IT IN YOUR backend.py
# MAKE SURE IT'S NOT INDENTED INSIDE ANOTHER FUNCTION!

# Updated backend.py - Replace the existing secure_interview_page route

# Add these modifications to your backend.py

# 1. Update the secure_interview_page function to handle reconnections
@app.route('/secure-interview/<token>', methods=['GET'])
def secure_interview_page(token):
    """Serve the secure interview page with proper session data and reconnection support"""
    try:
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(interview_token=token).first()
            if not candidate:
                return create_error_page(token, "Interview not found"), 404

            # Optional expiry extension logic omitted for brevity
            is_reconnection = bool(getattr(candidate, 'interview_started_at', None))

            company_name = os.getenv('COMPANY_NAME', 'Our Company')
            if getattr(candidate, 'company_name', None):
                company_name = candidate.company_name

            job_description = getattr(candidate, 'job_description', f'Interview for {candidate.job_title} position')

            interview_data = {
                'token': token,
                'candidateId': candidate.id,
                'candidateName': candidate.name,
                'candidateEmail': candidate.email,
                'position': candidate.job_title,
                'company': company_name,
                'knowledgeBaseId': getattr(candidate, 'knowledge_base_id', None),
                'sessionId': getattr(candidate, 'interview_session_id', None),
                'status': 'active',
                'jobDescription': job_description,
                'atsScore': candidate.ats_score,
                'resumePath': candidate.resume_path,
                'isReconnection': is_reconnection,
                'previousSessionData': {
                    'questionsAsked': getattr(candidate, 'interview_total_questions', 0),
                    'questionsAnswered': getattr(candidate, 'interview_answered_questions', 0),
                    'duration': getattr(candidate, 'interview_duration', 0)
                }
            }

            return create_interview_landing_page(interview_data, token), 200
        finally:
            session.close()
    except Exception as e:
        logger.exception("Error in interview route")
        return create_error_page(token, str(e)), 500


# 2. Add a new endpoint to validate and refresh interview tokens
@app.route('/api/interview/validate-token/<token>', methods=['GET', 'POST'])
def validate_interview_token(token):
    """Validate interview token and extend expiration if needed"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"valid": False, "error": "Invalid token"}), 404
        
        # Check expiration
        is_expired = False
        if hasattr(candidate, 'interview_expires_at') and candidate.interview_expires_at:
            if candidate.interview_expires_at < datetime.now():
                # Allow grace period of 24 hours
                time_since_expiry = datetime.now() - candidate.interview_expires_at
                if time_since_expiry.total_seconds() > 86400:
                    is_expired = True
                else:
                    # Extend expiration
                    candidate.interview_expires_at = datetime.now() + timedelta(days=7)
                    session.commit()
        
        # Get session info
        session_info = {
            "valid": not is_expired,
            "candidate_name": candidate.name,
            "position": candidate.job_title,
            "interview_started": bool(getattr(candidate, 'interview_started_at', None)),
            "interview_completed": bool(getattr(candidate, 'interview_completed_at', None)),
            "knowledge_base_id": getattr(candidate, 'knowledge_base_id', None),
            "can_reconnect": True,
            "questions_asked": getattr(candidate, 'interview_total_questions', 0),
            "questions_answered": getattr(candidate, 'interview_answered_questions', 0)
        }
        
        if request.method == 'POST':
            # Refresh session
            candidate.last_accessed = datetime.now()
            session.commit()
        
        return jsonify(session_info), 200
        
    except Exception as e:
        logger.error(f"Error validating token: {e}")
        return jsonify({"valid": False, "error": str(e)}), 500
    finally:
        session.close()


# 3. Update the create_interview_landing_page to show reconnection status
def create_interview_landing_page(interview_data, token):
    """Create the landing page HTML with reconnection support without illegal f-strings."""
    import json
    interview_json = json.dumps(interview_data)

    # Pre-build dynamic fragments (NO backslashes in f-string expressions):
    is_reconnection = bool(interview_data.get('isReconnection'))
    prev = interview_data.get('previousSessionData') or {}
    reconnect_block = ""
    if is_reconnection:
        reconnect_block = (
            "<div class=\"reconnect-info\">"
            "<h3>🔄 Welcome Back!</h3>"
            "<p>You're reconnecting to your interview session.</p>"
            f"<p><strong>Questions asked:</strong> {prev.get('questionsAsked', 0)}</p>"
            f"<p><strong>Questions answered:</strong> {prev.get('questionsAnswered', 0)}</p>"
            "</div>"
        )
    session_type_label = "Reconnection" if is_reconnection else "New Session"
    continue_or_start = "Continue" if is_reconnection else "Start"
    continue_lower = "continue" if is_reconnection else "start"
    progress_line = (
        f"<p><strong>📊 Progress:</strong> {prev.get('questionsAnswered', 0)} questions completed</p>"
        if is_reconnection else ""
    )

    # Build the HTML
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    candidate_name = interview_data['candidateName']
    position = interview_data['position']
    company = interview_data['company']
    knowledge_base_id = interview_data.get('knowledgeBaseId', '')
    candidate_id = interview_data['candidateId']

    return f"""<!DOCTYPE html>
<html>
<head>
  <title>AI Interview - {position}</title>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <style>
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
      margin: 0; padding: 0;
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      min-height: 100vh; display: flex; align-items: center; justify-content: center;
    }}
    .container {{
      background: white; padding: 2rem; border-radius: 15px;
      box-shadow: 0 20px 40px rgba(0,0,0,0.1);
      text-align: center; max-width: 600px; width: 90%;
    }}
    .header {{ color: #333; margin-bottom: 1.5rem; }}
    .info-box {{
      background: #f8f9fa; padding: 1.5rem; border-radius: 10px;
      margin: 1rem 0; border-left: 5px solid #667eea;
    }}
    .reconnect-info {{
      background: #e3f2fd; padding: 1rem; border-radius: 8px;
      margin: 1rem 0; border-left: 5px solid #2196f3;
    }}
    .success-badge {{ color: #28a745; font-weight: bold; font-size: 1.1em; }}
    .start-btn {{
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      color: white; padding: 15px 30px; border: none; border-radius: 50px;
      cursor: pointer; font-size: 18px; font-weight: bold; text-decoration: none;
      display: inline-block; margin: 20px 10px; transition: transform 0.3s, box-shadow 0.3s;
    }}
    .start-btn:hover {{ transform: translateY(-2px); box-shadow: 0 10px 25px rgba(102,126,234,0.3); }}
    .instructions {{ text-align: left; margin: 1.5rem 0; padding: 1.5rem; background: #e3f2fd; border-radius: 10px; }}
    .debug-info {{ background: #fff3cd; padding: 1rem; border-radius: 8px; font-size: 14px; margin-top: 15px; text-align: left; }}
  </style>
</head>
<body>
  <div class="container">
    <h1 class="header">🤖 AI Interview Portal</h1>

    {reconnect_block}

    <div class="info-box">
      <p><strong>👤 Candidate:</strong> {candidate_name}</p>
      <p><strong>💼 Position:</strong> {position}</p>
      <p><strong>🏢 Company:</strong> {company}</p>
      <p class="success-badge">✅ Interview Link Active & Ready</p>
      <p><strong>🔄 Session Type:</strong> {session_type_label}</p>
    </div>

    <div class="instructions">
      <h3>📋 Before {continue_or_start} Your Interview:</h3>
      <ul>
        <li><strong>Internet:</strong> Ensure stable connection</li>
        <li><strong>Camera & Mic:</strong> Test and allow permissions</li>
        <li><strong>Environment:</strong> Find a quiet, well-lit space</li>
        <li><strong>Materials:</strong> Have your resume ready</li>
      </ul>
    </div>

    <div style="margin: 25px 0;">
      <p><strong>⏱️ Duration:</strong> 30-45 minutes</p>
      <p><strong>🎥 Format:</strong> AI-powered video interview</p>
      {progress_line}
    </div>

    <button onclick="startInterview()" class="start-btn">🚀 {continue_or_start} AI Interview</button>

    <div class="debug-info">
      <strong>🔧 System Status:</strong><br/>
      Token: {token}<br/>
      Knowledge Base: {knowledge_base_id}<br/>
      Candidate ID: {candidate_id}<br/>
      Status: Ready ✅<br/>
      Session Type: {session_type_label}<br/>
      Time: {now_str}
    </div>

    <div style="margin-top: 30px; font-size: 12px; color: #666;">
      <p>💡 Need help? Contact our support team</p>
      <p>🕐 Interview link valid for multiple sessions</p>
    </div>
  </div>

  <script>
    const interviewData = {interview_json};
    function startInterview() {{
      console.log('🚀 Starting interview with data:', interviewData);
      sessionStorage.setItem('interviewData', JSON.stringify(interviewData));
      fetch('/api/avatar/interview/{token}', {{
        method: 'POST',
        headers: {{ 'Content-Type': 'application/json' }},
        body: JSON.stringify({{ action: '{continue_lower}' }})
      }}).catch(() => {{ /* best effort */ }});
      window.location.href = 'http://localhost:3001/interview/{token}';
    }}

    fetch('/api/interview/validate-token/{token}', {{ method: 'POST' }})
      .then(r => r.json())
      .then(d => {{
        if (!d.valid) {{
          alert('This interview link has expired. Please contact HR for assistance.');
          document.querySelector('.start-btn').disabled = true;
        }}
      }})
      .catch(() => {{ /* noop */ }});
  </script>
</body>
</html>"""


def create_expired_interview_page(token):
    """Create page for expired interviews"""
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview Expired</title>
  <style>
    body {{ font-family: Arial, sans-serif; text-align: center; margin-top: 100px; background: #f8f9fa; }}
    .container {{ max-width: 500px; margin: 0 auto; padding: 2rem; background: white; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;">⏰ Interview Link Expired</h1>
    <p>This interview link has expired. Please contact HR for a new interview link.</p>
    <p><strong>Token:</strong> {token}</p>
  </div>
</body>
</html>""", 410


def create_error_page(token, error):
    """Create error page"""
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview Error</title>
  <style>
    body {{ font-family: Arial, sans-serif; text-align: center; margin-top: 100px; background: #f8f9fa; }}
    .container {{ max-width: 500px; margin: 0 auto; padding: 2rem; background: white; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;">🔥 Interview System Error</h1>
    <p>There was an error loading the interview.</p>
    <p><strong>Error:</strong> {error}</p>
    <p><strong>Token:</strong> {token}</p>
    <p>The system has been notified. Please try again or contact support.</p>
    <button onclick="window.location.reload()" style="padding:10px 20px; background:#007bff; color:#fff; border:none; border-radius:5px; cursor:pointer;">
      Retry
    </button>
  </div>
</body>
</html>"""

@app.route('/api/scrape_assessment_results', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=3, time_window=300)  # Max 3 scraping requests per 5 minutes
def api_scrape_assessment_results():
    """API endpoint to scrape assessment results for a specific assessment"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        data = request.json
        assessment_name = data.get('assessment_name')
        
        if not assessment_name:
            return jsonify({"success": False, "message": "assessment_name is required"}), 400
        
        logger.info(f"Starting results scraping for assessment: {assessment_name}")
        
        # Start scraping in a separate thread
        scraping_thread = threading.Thread(
            target=lambda: run_scraping_with_monitoring(assessment_name),
            daemon=True,
            name=f"scraping_{assessment_name.replace(' ', '_')}_{int(time.time())}"
        )
        scraping_thread.start()
        
        return jsonify({
            "success": True,
            "message": f"Started scraping results for '{assessment_name}'",
            "estimated_time": "2-5 minutes"
        }), 200
        
    except Exception as e:
        logger.error(f"Error in scrape_assessment_results: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/scrape_all_pending_results', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=1, time_window=600)  # Max 1 bulk scraping per 10 minutes
def api_scrape_all_pending_results():
    """API endpoint to scrape all pending assessment results"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        logger.info("Starting bulk results scraping for all pending assessments")
        
        # Start bulk scraping in a separate thread
        scraping_thread = threading.Thread(
            target=lambda: run_bulk_scraping_with_monitoring(),
            daemon=True,
            name=f"bulk_scraping_{int(time.time())}"
        )
        scraping_thread.start()
        
        return jsonify({
            "success": True,
            "message": "Started bulk scraping for all pending assessments",
            "estimated_time": "5-15 minutes"
        }), 200
        
    except Exception as e:
        logger.error(f"Error in scrape_all_pending_results: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/scrape_assessment_results', methods=['GET','OPTIONS'])
def api_scraping_status():
    """Get status of running scraping operations"""
    try:
        # Get active scraping threads
        active_threads = []
        for thread in threading.enumerate():
            if thread.name.startswith(('scraping_', 'bulk_scraping_')):
                thread_info = {
                    "name": thread.name,
                    "is_alive": thread.is_alive(),
                    "daemon": thread.daemon
                }
                active_threads.append(thread_info)
        
        return jsonify({
            "success": True,
            "active_operations": len(active_threads),
            "operations": active_threads
        }), 200
        
    except Exception as e:
        logger.error(f"Error in scraping_status: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500
def run_scraping_with_monitoring(assessment_name: str):
    """Wrapper to run scraping with monitoring and error handling"""
    start_time = time.time()
    
    try:
        logger.info(f"Starting monitored scraping for assessment: {assessment_name}")
        
        # Import and run the scraping function
        try:
            from testlify_results_scraper import scrape_assessment_results_by_name
        except ImportError as e:
            logger.error(f"Failed to import scraper: {e}")
            notify_admin(
                "Scraper Import Error",
                f"Could not import results scraper: {str(e)}. Please ensure testlify_results_scraper.py is available."
            )
            return
        
        # Run the async scraping function
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            results = loop.run_until_complete(scrape_assessment_results_by_name(assessment_name))
        finally:
            loop.close()
        
        duration = time.time() - start_time
        logger.info(f"Scraping completed successfully in {duration:.2f} seconds. Found {len(results)} candidates.")
        
        # Send success notification
        notify_admin(
            "Assessment Results Scraping Completed",
            f"Assessment: {assessment_name}\nCandidates processed: {len(results)}\nDuration: {duration:.2f} seconds"
        )
        
    except Exception as e:
        duration = time.time() - start_time
        error_msg = f"Scraping failed for assessment '{assessment_name}' after {duration:.2f} seconds"
        logger.error(error_msg, exc_info=True)
        
        # Send failure notification
        notify_admin(
            "Assessment Results Scraping Failed",
            error_msg,
            error_details=traceback.format_exc()
        )

@app.route('/api/get-interview/<token>', methods=['GET', 'POST'])
def get_interview(token):
    from datetime import timezone  # ensure timezone available here
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        if not candidate:
            return jsonify({'error': 'Interview not found'}), 404

        if request.method == 'POST':
            body = request.get_json(silent=True) or {}
            action = body.get('action')

            if action == 'start':
                candidate.interview_started_at = datetime.now(timezone.utc)
                session.commit()
                return jsonify({"success": True}), 200

            elif action == 'complete':
                candidate.interview_completed_at = datetime.now(timezone.utc)
                transcript = body.get('transcript')
                if transcript:
                    candidate.interview_transcript = transcript
                session.commit()
                return jsonify({"success": True}), 200

        # unify KB id across both possible columns
        kb = getattr(candidate, 'knowledge_base_id', None) or getattr(candidate, 'interview_kb_id', None)
        position = getattr(candidate, "job_title", None) or getattr(candidate, "position", "") or "Interview"

        data = {
            "id": getattr(candidate, "id", None),
            "token": getattr(candidate, "interview_token", None),
            "candidateId": getattr(candidate, "id", None),
            "candidateName": getattr(candidate, "name", "") or "",
            "candidateEmail": getattr(candidate, "email", "") or "",
            "position": position,
            "company": getattr(candidate, "company_name", None)
                       or getattr(candidate, "company", None)
                       or os.getenv("COMPANY_NAME", "Our Company"),
            "knowledgeBaseId": kb,
            "status": "active" if getattr(candidate, "interview_scheduled", False) else "inactive",
            "jobDescription": getattr(candidate, "job_description", None)
                              or f"Interview for {position} position",
            "resumeLink": getattr(candidate, "resume_link", None) or getattr(candidate, "resume_path", None),
            "createdAt": getattr(candidate, "interview_created_at", None).isoformat()
                         if getattr(candidate, "interview_created_at", None) else None,
        }
        return jsonify(data), 200

    except Exception as e:
        logger.error(f"Error in get_interview: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500
    finally:
        session.close()


# --- TEMP DEBUG: check KB columns by interview token ---
@app.route('/api/debug/kb/<token>', methods=['GET'])
def debug_kb(token):
    s = SessionLocal()
    try:
        c = s.query(Candidate).filter_by(interview_token=token).first()
        if not c:
            return jsonify({"error": "not found", "token": token}), 404
        return jsonify({
            "id": c.id,
            "name": c.name,
            "email": c.email,
            "interview_token": c.interview_token,
            "knowledge_base_id": getattr(c, "knowledge_base_id", None),
            "interview_kb_id": getattr(c, "interview_kb_id", None),
        }), 200
    except Exception as e:
        logger.exception("debug_kb error")
        return jsonify({"error": str(e)}), 500
    finally:
        s.close()


@app.route('/api/verify-knowledge-base/<candidate_id>', methods=['GET'])
def verify_knowledge_base(candidate_id):
    """Verify knowledge base content for debugging"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Extract resume content
        resume_content = ""
        if candidate.resume_path and os.path.exists(candidate.resume_path):
            resume_content = extract_resume_content(candidate.resume_path)
        
        # Build preview WITHOUT backslashes in f-string expressions
        questions = generate_interview_questions(
            candidate_name=candidate.name,
            position=candidate.job_title,
            resume_content=resume_content,
            job_description=getattr(candidate, 'job_description', f"Position: {candidate.job_title}")
        )
        questions_preview = (questions[:1000] + "...") if len(questions) > 1000 else questions
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "position": candidate.job_title,
            "knowledge_base_id": getattr(candidate, "knowledge_base_id", None),
            "resume_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "resume_content_length": len(resume_content),
            "interview_scheduled": candidate.interview_scheduled,
            "generated_questions_preview": questions_preview,
            "total_questions_length": len(questions)
        }), 200
    finally:
        session.close()


def extract_skills_from_resume(resume_content):
    """Extract technical skills from resume"""
    skills = []
    resume_lower = resume_content.lower()
    tech_skills = [
        'python', 'javascript', 'java', 'c++', 'c#', 'react', 'angular', 'vue',
        'node.js', 'django', 'flask', 'spring', 'sql', 'nosql', 'mongodb',
        'postgresql', 'mysql', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
        'git', 'ci/cd', 'machine learning', 'data science', 'api', 'rest',
        'graphql', 'typescript', 'golang', 'rust', 'swift', 'kotlin'
    ]
    for skill in tech_skills:
        if skill in resume_lower:
            skills.append(skill.title())
    return skills[:10]


def extract_projects_from_resume(resume_content):
    """Extract project names from resume"""
    projects = []
    lines = resume_content.split('\n')
    for i, line in enumerate(lines):
        if 'project' in line.lower():
            if ':' in line:
                project_name = line.split(':', 1)[1].strip()[:50]
                if project_name:
                    projects.append(project_name)
    return projects[:5]


def extract_experience_years(resume_content):
    """Extract years of experience from resume"""
    import re
    pattern = r'(\d+)\+?\s*years?\s*(?:of\s*)?experience'
    match = re.search(pattern, resume_content.lower())
    if match:
        return f"{match.group(1)}+ years"
    year_pattern = r'20\d{2}'
    years = re.findall(year_pattern, resume_content)
    if len(years) >= 2:
        min_year = min(int(y) for y in years)
        max_year = max(int(y) for y in years)
        experience = max_year - min_year
        if experience > 0:
            return f"{experience}+ years"
    return "Not specified"


@app.route('/api/avatar/get-access-token', methods=['POST', 'OPTIONS'])
def get_avatar_access_token():
    """Generate HeyGen access token for avatar session"""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        heygen_key = os.getenv('HEYGEN_API_KEY', 'your_heygen_api_key_here')
        return heygen_key, 200, {'Content-Type': 'text/plain'}
    except Exception as e:
        logger.error(f"Error getting access token: {e}")
        return jsonify({"error": "Failed to get access token"}), 500


@app.route('/api/debug-schedule-interview', methods=['POST'])
def debug_schedule_interview():
    """Debug: try to create a KB and echo responses for troubleshooting"""
    try:
        data = request.json or {}
        api_key = os.getenv('HEYGEN_API_KEY')
        if not api_key:
            return jsonify({
                "error": "HEYGEN_API_KEY not found in environment variables",
                "fix": "Add HEYGEN_API_KEY to your .env file"
            }), 400
        if len(api_key) < 20:
            return jsonify({"error": "HEYGEN_API_KEY seems too short", "length": len(api_key)}), 400

        test_payload = {
            'name': f'Test_Interview_{int(time.time())}',
            'description': 'Test knowledge base',
            'content': 'Test interview questions: 1. Tell me about yourself. 2. Why this role?',
            'opening_line': 'Hello, this is a test interview.'
        }
        try:
            heygen_response = requests.post(
                'https://api.heygen.com/v1/streaming/knowledge_base',
                headers={
                    'X-Api-Key': api_key,
                    'Content-Type': 'application/json',
                    'Accept': 'application/json'
                },
                json=test_payload,
                timeout=30
            )
            if heygen_response.ok:
                kb_data = heygen_response.json()
                return jsonify({
                    "success": True,
                    "knowledge_base_id": (kb_data.get('data', {}) or {}).get('knowledge_base_id') or
                                         (kb_data.get('data', {}) or {}).get('id') or
                                         kb_data.get('knowledge_base_id') or kb_data.get('id'),
                    "full_response": kb_data
                }), 200
            else:
                return jsonify({
                    "success": False,
                    "error": "HeyGen API error",
                    "status_code": heygen_response.status_code,
                    "body": heygen_response.text[:500]
                }), 400
        except requests.exceptions.RequestException as e:
            return jsonify({"success": False, "error": "Request failed", "details": str(e)}), 500
    except Exception as e:
        return jsonify({
            "success": False,
            "error": "Unexpected error",
            "details": str(e),
            "traceback": traceback.format_exc()
        }), 500


# Add this enhanced function to your backend.py

def create_structured_interview_kb(candidate_name, position, company, resume_content, job_description):
    """Create a highly structured knowledge base for professional interviews"""
    
    # Extract key information
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    experience = extract_experience_years(resume_content) if resume_content else "Not specified"
    projects = extract_projects_from_resume(resume_content) if resume_content else []
    
    # Build the structured interview content
    structured_content = f"""
STRICT INTERVIEW PROTOCOL - FOLLOW EXACTLY:

YOU ARE: A professional AI interviewer conducting a structured technical interview.
YOUR BEHAVIOR: Professional, clear, structured. NO casual conversation.

CANDIDATE DETAILS:
- Name: {candidate_name}
- Position: {position}
- Company: {company}
- Experience: {experience}
- Key Skills: {', '.join(skills[:5]) if skills else 'General skills'}

INTERVIEW STRUCTURE - ASK THESE QUESTIONS IN EXACT ORDER:

=== QUESTION 1: SELF INTRODUCTION (ALWAYS ASK FIRST) ===
"Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background and what led you to apply for this role?"

=== QUESTION 2: TECHNICAL EXPERIENCE ===
{"I see from your resume that you have experience with " + skills[0] + ". Can you tell me about a specific project where you used " + skills[0] + " and what challenges you faced?" if skills else "Can you tell me about your most significant technical project and the technologies you used?"}

=== QUESTION 3: PROBLEM SOLVING ===
"That's interesting. Now, let me ask you about problem-solving. Can you describe a time when you encountered a complex technical problem and walk me through how you approached and solved it?"

=== QUESTION 4: TEAMWORK ===
"Great. Let's talk about teamwork. Tell me about a time when you had to collaborate with other team members on a challenging project. How did you handle any conflicts or disagreements?"

=== QUESTION 5: SPECIFIC SKILL DEEP DIVE ===
{"I noticed you also have experience with " + skills[1] + ". What's the most complex thing you've built using " + skills[1] + "? Please be specific about the technical details." if len(skills) > 1 else "What would you say is your strongest technical skill, and can you give me a detailed example of how you've applied it?"}

=== QUESTION 6: LEARNING ABILITY ===
"Technology evolves rapidly. Can you tell me about a time when you had to quickly learn a new technology or framework for a project? How did you approach the learning process?"

=== QUESTION 7: ROLE-SPECIFIC ===
"Let's talk specifically about this {position} role. Based on your understanding of the position, how do you see your skills and experience contributing to our team in the first 90 days?"

=== QUESTION 8: CHALLENGES ===
"What do you think would be the biggest challenge for you in this role, and how would you address it?"

=== QUESTION 9: CAREER GOALS ===
"Where do you see your career heading in the next 3-5 years, and how does this {position} role fit into those plans?"

=== QUESTION 10: CANDIDATE QUESTIONS ===
"Thank you for your answers. Now, do you have any questions for me about the role, the team, or {company}?"

CRITICAL RULES:
1. When you receive "INIT_INTERVIEW", IMMEDIATELY ask Question 1
2. Ask ONE question at a time
3. Wait for complete answer before next question
4. After each answer say: "Thank you for sharing that. [Next question]"
5. Stay professional - NO casual chat
6. If candidate says just "Hello", respond with Question 1
7. Track which question you're on to avoid repetition

FORBIDDEN BEHAVIORS:
- Do NOT say: "hey", "cool", "chat", "Oh", or use casual language
- Do NOT have conversations outside these 10 questions
- Do NOT ask random questions
- Do NOT give short one-word responses

RESUME CONTEXT:
{resume_content[:2000] if resume_content else 'No resume provided'}

JOB REQUIREMENTS:
{job_description[:1000] if job_description else f'Standard requirements for {position}'}

Remember: You are conducting a PROFESSIONAL STRUCTURED INTERVIEW. Stay on script!
"""
    
    return structured_content


@app.route('/api/create-knowledge-base', methods=['POST', 'OPTIONS'])
def create_knowledge_base_enhanced():
    """Create a structured interview knowledge base"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json or {}
        candidate_name = data.get('candidateName')
        position = data.get('position')
        company = data.get('company', 'Our Company')
        token = data.get('token')
        
        if not candidate_name or not position:
            return jsonify({"error": "candidateName and position are required"}), 400
        
        logger.info(f"🧠 Creating structured KB for: {candidate_name} - {position}")
        
        # Get resume content
        resume_content = ""
        candidate_id = None
        job_description = data.get('jobDescription', '')
        
        if token:
            session = SessionLocal()
            try:
                cand = session.query(Candidate).filter_by(interview_token=token).first()
                if cand:
                    candidate_id = cand.id
                    if cand.resume_path and os.path.exists(cand.resume_path):
                        resume_content = extract_resume_content(cand.resume_path)
                    job_description = job_description or getattr(cand, 'job_description', '')
            finally:
                session.close()
        
        # Create structured interview content
        kb_content = create_structured_interview_kb(
            candidate_name=candidate_name,
            position=position,
            company=company,
            resume_content=resume_content,
            job_description=job_description
        )
        
        # Create opening line that immediately asks first question
        opening_line = f"Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background and what led you to apply for this role?"
        
        heygen_key = os.getenv('HEYGEN_API_KEY')
        if not heygen_key:
            # Fallback
            fallback_kb_id = f"kb_structured_{candidate_name.replace(' ', '_')}_{int(time.time())}"
            logger.warning(f"No HeyGen API key, using fallback: {fallback_kb_id}")
            
            if candidate_id:
                session = SessionLocal()
                try:
                    cand = session.query(Candidate).filter_by(id=candidate_id).first()
                    if cand:
                        cand.knowledge_base_id = fallback_kb_id
                        cand.interview_kb_content = kb_content
                        session.commit()
                finally:
                    session.close()
            
            return jsonify({
                "success": True,
                "knowledgeBaseId": fallback_kb_id,
                "fallback": True
            }), 200
        
        # Create HeyGen knowledge base with strict instructions
        heygen_payload = {
            "name": f"Structured_Interview_{candidate_name.replace(' ', '_')}_{int(time.time())}",
            "description": f"Structured technical interview for {candidate_name} - {position}",
            "content": kb_content,
            "opening_line": opening_line,
            "custom_prompt": "You are a professional interviewer. Follow the structured questions EXACTLY as provided. No casual chat. When you receive INIT_INTERVIEW, immediately ask the first question about self-introduction.",
            "prompt": kb_content,  # Some endpoints use 'prompt' instead of 'content'
            "voice_settings": {
                "rate": 1.0,
                "emotion": "professional"
            }
        }
        
        # Try multiple endpoints
        endpoints = [
            "https://api.heygen.com/v1/streaming/knowledge_base/create",
            "https://api.heygen.com/v1/streaming/knowledge_base",
            "https://api.heygen.com/v1/streaming_avatar/knowledge_base"
        ]
        
        kb_id = None
        successful_endpoint = None
        
        for endpoint in endpoints:
            try:
                logger.info(f"Trying endpoint: {endpoint}")
                resp = requests.post(
                    endpoint,
                    headers={
                        "X-Api-Key": heygen_key,
                        "Content-Type": "application/json",
                        "Accept": "application/json"
                    },
                    json=heygen_payload,
                    timeout=30
                )
                
                if resp.ok:
                    data = resp.json()
                    kb_id = (
                        data.get('data', {}).get('knowledge_base_id') or
                        data.get('knowledge_base_id') or
                        data.get('id')
                    )
                    if kb_id:
                        successful_endpoint = endpoint
                        logger.info(f"✅ KB created successfully: {kb_id}")
                        break
                else:
                    logger.error(f"Endpoint {endpoint} failed: {resp.status_code} - {resp.text[:200]}")
                    
            except Exception as e:
                logger.error(f"Error with endpoint {endpoint}: {e}")
        
        if not kb_id:
            kb_id = f"kb_structured_{candidate_name.replace(' ', '_')}_{int(time.time())}"
            logger.warning(f"All endpoints failed, using fallback: {kb_id}")
        
        # Save to database
        if candidate_id:
            session = SessionLocal()
            try:
                cand = session.query(Candidate).filter_by(id=candidate_id).first()
                if cand:
                    cand.knowledge_base_id = kb_id
                    cand.interview_kb_content = kb_content
                    session.commit()
            finally:
                session.close()
        
        return jsonify({
            "success": True,
            "knowledgeBaseId": kb_id,
            "endpoint_used": successful_endpoint,
            "structured": True,
            "question_count": 10
        }), 200
        
    except Exception as e:
        logger.error(f"KB creation error: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

def create_enhanced_kb_content(candidate_name, position, company, resume_content):
    """Create enhanced knowledge base content WITHOUT putting backslashes inside f-string expressions."""
    # Build any strings that contain backslashes/newlines first:
    resume_highlights = (
        f"Resume Highlights:\n{resume_content[:2000]}..."
        if resume_content else
        "No resume content available - focus on standard interview questions"
    )
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    experience_years = extract_experience_years(resume_content) if resume_content else "Not specified"

    skills_line = ", ".join(skills) if skills else "General software engineering skills"
    # Now the f-string only injects already-built variables (safe):
    return (
        "INTERVIEW SYSTEM CONFIGURATION\n"
        "==============================\n"
        f"🎯 MISSION: Conduct a professional, comprehensive interview for {candidate_name}\n"
        f"📋 POSITION: {position}\n"
        f"🏢 COMPANY: {company}\n"
        "⏱️ DURATION: 30-45 minutes\n"
        "🤖 MODE: AI-Powered Structured Interview\n\n"
        "CANDIDATE BACKGROUND\n"
        "====================\n"
        f"{resume_highlights}\n"
        f"Experience: {experience_years}\n"
        f"Key Skills: {skills_line}\n\n"
        "INTERVIEW FLOW\n"
        "==============\n"
        "1) Intro & Warm-up\n"
        "2) Skills Deep Dive\n"
        "3) Problem Solving\n"
        "4) Behavioral\n"
        "5) Wrap-up\n"
    )

# backend.py - Add this debug endpoint
@app.route('/api/verify-interview-system/<token>', methods=['GET'])
def verify_interview_system(token):
    """Verify complete interview system setup"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Check all components
        checks = {
            "candidate_found": True,
            "resume_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "knowledge_base_created": bool(candidate.knowledge_base_id),
            "heygen_api_configured": bool(os.getenv('HEYGEN_API_KEY')),
            "session_configured": bool(candidate.interview_session_id),
            "recording_ready": True,
            "qa_tracking_ready": True
        }
        
        # Extract sample questions
        sample_questions = []
        if candidate.resume_path:
            resume_content = extract_resume_content(candidate.resume_path)
            skills = extract_skills_from_resume(resume_content)
            sample_questions = [
                f"Tell me about your experience with {skills[0]}" if skills else "Tell me about yourself",
                "What interests you about this position?",
                "Describe a challenging project you've worked on"
            ]
        
        return jsonify({
            "status": "ready" if all(checks.values()) else "issues_found",
            "checks": checks,
            "candidate_info": {
                "name": candidate.name,
                "position": candidate.job_title,
                "knowledge_base_id": candidate.knowledge_base_id
            },
            "sample_questions": sample_questions,
            "recommendations": [
                "Ensure HEYGEN_API_KEY is set" if not checks["heygen_api_configured"] else None,
                "Upload candidate resume" if not checks["resume_exists"] else None,
                "Create knowledge base" if not checks["knowledge_base_created"] else None
            ]
        }), 200
        
    finally:
        session.close()
@app.route('/api/verify-knowledge-base', methods=['GET'])
def verify_kb_by_query():
    kb_id = request.args.get('id')
    if not kb_id:
        return jsonify({"error": "id is required"}), 400
    # If you want: actually check existence in your DB or HeyGen here
    return jsonify({"ok": True, "knowledge_base_id": kb_id}), 200


@app.route('/api/avatar/interviews', methods=['POST','OPTIONS'], endpoint='save_interview_v2')
def save_interview_v2():
    """Create/refresh an interview token for a candidate and persist expiry."""
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.status_code = 200
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        return resp

    from sqlalchemy.exc import SQLAlchemyError
    try:
        data = request.get_json(silent=True) or {}
        candidate_email = data.get('candidateEmail')
        incoming_kb_id = data.get('knowledgeBaseId')

        if not candidate_email:
            return jsonify({"error": "candidateEmail is required"}), 400

        interview_token = str(uuid.uuid4())
        expires_at = datetime.now(timezone.utc) + timedelta(days=7)

        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(email=candidate_email).first()
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404

            candidate.interview_token = interview_token
            candidate.interview_expires_at = expires_at
            if incoming_kb_id:
                # only overwrite if caller explicitly sends
                candidate.knowledge_base_id = incoming_kb_id

            session.commit()

            return jsonify({
                "token": interview_token,
                "expiresAt": expires_at.isoformat(),
                "knowledgeBaseId": getattr(candidate, "knowledge_base_id", None)
            }), 200

        except SQLAlchemyError:
            session.rollback()
            logger.exception("DB error saving interview")
            return jsonify({"error": "Database error saving interview"}), 500
        finally:
            session.close()

    except Exception:
        logger.exception("Error saving interview")
        return jsonify({"error": "Failed to save interview"}), 500

@app.route('/api/debug/find-candidate', methods=['POST'])
def debug_find_candidate():
    from sqlalchemy.exc import SQLAlchemyError
    data = request.get_json(silent=True) or {}
    email = data.get('candidateEmail')
    if not email:
        return jsonify({"error": "candidateEmail required"}), 400
    s = SessionLocal()
    try:
        c = s.query(Candidate).filter_by(email=email).first()
        if not c:
            return jsonify({"found": False, "email": email}), 404
        return jsonify({
            "found": True,
            "id": c.id,
            "name": c.name,
            "email": c.email,
            "has_interview_token": hasattr(c, "interview_token"),
            "has_interview_expires_at": hasattr(c, "interview_expires_at"),
            "has_knowledge_base_id": hasattr(c, "knowledge_base_id"),
            "current_interview_token": getattr(c, "interview_token", None),
            "current_kb": getattr(c, "knowledge_base_id", None)
        }), 200
    except SQLAlchemyError as e:
        return jsonify({"error": "db", "details": str(e)}), 500
    finally:
        s.close()

@app.route('/api/avatar/interviews', methods=['POST','OPTIONS'])
def save_interview():
    if request.method == 'OPTIONS':
        return ('', 200)

    data = request.get_json(silent=True) or {}
    email = data.get('candidateEmail')
    if not email:
        return jsonify({"error":"candidateEmail is required"}), 400

    interview_token = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)

    s = SessionLocal()
    try:
        c = s.query(Candidate).filter_by(email=email).first()
        if not c:
            return jsonify({"error":"Candidate not found","email":email}), 404

        # Write fields (these must exist in your model/DB)
        c.interview_token = interview_token
        c.interview_expires_at = expires_at

        incoming_kb = data.get('knowledgeBaseId')
        if incoming_kb:
            c.knowledge_base_id = incoming_kb

        s.commit()
        return jsonify({
            "token": interview_token,
            "expiresAt": expires_at.isoformat(),
            "knowledgeBaseId": getattr(c,"knowledge_base_id",None),
            "candidateEmail": email
        }), 200

    except SQLAlchemyError as e:
        s.rollback()
        # TEMP: return the DB error so we can see it quickly
        return jsonify({"error":"db","details":str(e)}), 500
    except Exception as e:
        s.rollback()
        return jsonify({"error":"server","details":str(e)}), 500
    finally:
        s.close()


@app.route('/api/avatar/interview/<token>', methods=['POST'])
def api_avatar_interview(token):
    """Handle avatar interview updates"""
    try:
        data = request.json
        action = data.get('action')
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(interview_token=token).first()
            
            if not candidate:
                return jsonify({"error": "Interview not found"}), 404
            
            if action == 'start':
                candidate.interview_started_at = datetime.now()
                message = "Interview started"
            elif action == 'complete':
                candidate.interview_completed_at = datetime.now()
                if data.get('transcript'):
                    candidate.interview_transcript = json.dumps(data['transcript'])
                message = "Interview completed"
            else:
                message = "Interview updated"
            
            session.commit()
            
            return jsonify({
                "success": True,
                "message": message
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error in avatar interview {token}: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview-automation/status', methods=['GET', 'OPTIONS'])
def get_automation_status():
    """Get interview automation system status"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        from interview_automation import interview_automation
        
        status = {
            'is_running': interview_automation.is_running,
            'check_interval_minutes': interview_automation.check_interval / 60,
            'next_check': 'Running' if interview_automation.is_running else 'Stopped'
        }
        
        # Get statistics
        session = SessionLocal()
        try:
            stats = {
                'candidates_pending_interview': session.query(Candidate).filter(
                    and_(
                        Candidate.exam_completed == True,
                        Candidate.exam_percentage >= 70,
                        Candidate.interview_scheduled == False
                    )
                ).count(),
                'interviews_scheduled': session.query(Candidate).filter(
                    Candidate.interview_scheduled == True
                ).count(),
                'interviews_completed': session.query(Candidate).filter(
                    Candidate.interview_completed_at.isnot(None)
                ).count()
            }
            status['statistics'] = stats
        finally:
            session.close()
        
        return jsonify(status), 200
        
    except Exception as e:
        logger.error(f"Error getting automation status: {e}")
        return jsonify({"error": str(e)}), 500
    
@app.route('/api/v1/streaming_new', methods=['POST', 'OPTIONS'])
def streaming_new():
    """Proxy requests to HeyGen streaming API"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json or {}
        
        # Get HeyGen API key
        heygen_key = os.getenv('HEYGEN_API_KEY')
        if not heygen_key:
            return jsonify({"error": "HeyGen API key not configured"}), 500
        
        # Use correct HeyGen streaming endpoint
        heygen_streaming_url = "https://api.heygen.com/v1/streaming.new"
        
        headers = {
            'Content-Type': 'application/json',
            'X-Api-Key': heygen_key,
            'Accept': 'application/json'
        }
        
        print(f"🔄 Forwarding request to HeyGen: {heygen_streaming_url}")
        print(f"📦 Request data: {data}")
        
        response = requests.post(
            heygen_streaming_url, 
            json=data, 
            headers=headers,
            timeout=30
        )
        
        print(f"📡 HeyGen response status: {response.status_code}")
        
        if response.ok:
            response_data = response.json()
            return jsonify(response_data), response.status_code
        else:
            error_text = response.text
            print(f"❌ HeyGen API error: {error_text}")
            return jsonify({
                "error": "HeyGen API error", 
                "details": error_text,
                "status": response.status_code
            }), response.status_code
            
    except requests.exceptions.Timeout:
        print("⏰ HeyGen API timeout")
        return jsonify({"error": "Request timeout"}), 504
    except requests.exceptions.ConnectionError:
        print("🔌 HeyGen API connection error")
        return jsonify({"error": "Connection error"}), 503
    except Exception as e:
        print(f"💥 Streaming proxy error: {e}")
        return jsonify({"error": "Avatar service unavailable", "details": str(e)}), 500

@app.route('/api/get-access-token', methods=['POST', 'OPTIONS'])
def get_access_token():
    """Get HeyGen streaming access token"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        heygen_key = os.getenv('HEYGEN_API_KEY')
        if not heygen_key:
            return jsonify({"error": "HeyGen API key not configured"}), 500
        
        # Call HeyGen token creation API
        response = requests.post(
            'https://api.heygen.com/v1/streaming.create_token',
            headers={
                'Accept': 'application/json',
                'Content-Type': 'application/json',
                'X-Api-Key': heygen_key,
            },
            timeout=10
        )
        
        if not response.ok:
            error_text = response.text
            print(f"❌ HeyGen token error: {error_text}")
            return jsonify({"error": "Failed to get token", "details": error_text}), response.status_code
        
        data = response.json()
        if not data.get('data', {}).get('token'):
            return jsonify({"error": "No token in response"}), 500
        
        token = data['data']['token']
        print(f"✅ Token obtained successfully")
        
        # Return as plain text
        return Response(token, mimetype='text/plain', status=200)
        
    except Exception as e:
        print(f"❌ Token error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/debug/avatar', methods=['GET'])
def debug_avatar():
    """Debug avatar configuration"""
    return jsonify({
        "heygen_key_configured": bool(os.getenv('HEYGEN_API_KEY')),
        "heygen_key_length": len(os.getenv('HEYGEN_API_KEY', '')),
        "cors_enabled": True,
        "endpoints": [
            "/api/get-access-token",
            "/api/v1/streaming_new",
            "/secure-interview/<token>"
        ]
    }), 200

@app.route('/api/interview/results', methods=['GET'])
def get_interview_results():
    """Get all interview results with filtering options"""
    session = SessionLocal()
    try:
        # Get query parameters
        position = request.args.get('position')
        status = request.args.get('status')
        date_from = request.args.get('date_from')
        date_to = request.args.get('date_to')
        
        # Build query
        query = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        )
        
        # Apply filters
        if position:
            query = query.filter(Candidate.job_title == position)
        
        if status:
            if status == 'completed':
                query = query.filter(Candidate.interview_completed_at.isnot(None))
            elif status == 'pending':
                query = query.filter(Candidate.interview_completed_at.is_(None))
        
        if date_from:
            query = query.filter(Candidate.interview_date >= date_from)
        
        if date_to:
            query = query.filter(Candidate.interview_date <= date_to)
        
        # Get results
        candidates = query.all()
        
        # Format results
        results = []
        for candidate in candidates:
            results.append({
                'id': candidate.id,
                'name': candidate.name,
                'email': candidate.email,
                'job_title': candidate.job_title,
                'interview_date': candidate.interview_date.isoformat() if candidate.interview_date else None,
                'interview_completed_at': candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                'interview_ai_score': candidate.interview_ai_score,
                'interview_ai_technical_score': candidate.interview_ai_technical_score,
                'interview_ai_communication_score': candidate.interview_ai_communication_score,
                'interview_ai_problem_solving_score': candidate.interview_ai_problem_solving_score,
                'interview_ai_cultural_fit_score': candidate.interview_ai_cultural_fit_score,
                'interview_ai_overall_feedback': candidate.interview_ai_overall_feedback,
                'interview_final_status': candidate.interview_final_status,
                'interview_recording_url': candidate.interview_recording_url,
                'interview_transcript': candidate.interview_transcript
            })
        
        return jsonify({
            'success': True,
            'results': results,
            'total': len(results)
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting interview results: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/stats', methods=['GET'])
def get_interview_stats():
    """Get interview statistics"""
    session = SessionLocal()
    try:
        # Get all interviewed candidates
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        # Calculate statistics
        total = len(candidates)
        completed = len([c for c in candidates if c.interview_completed_at])
        with_scores = len([c for c in candidates if c.interview_ai_score])
        passed = len([c for c in candidates if c.interview_ai_score and c.interview_ai_score >= 70])
        
        avg_score = 0
        if with_scores > 0:
            total_score = sum(c.interview_ai_score for c in candidates if c.interview_ai_score)
            avg_score = total_score / with_scores
        
        # Skills averages
        skills = {
            'technical': 0,
            'communication': 0,
            'problem_solving': 0,
            'cultural_fit': 0
        }
        
        for c in candidates:
            if c.interview_ai_technical_score:
                skills['technical'] += c.interview_ai_technical_score
            if c.interview_ai_communication_score:
                skills['communication'] += c.interview_ai_communication_score
            if c.interview_ai_problem_solving_score:
                skills['problem_solving'] += c.interview_ai_problem_solving_score
            if c.interview_ai_cultural_fit_score:
                skills['cultural_fit'] += c.interview_ai_cultural_fit_score
        
        # Calculate averages
        for skill in skills:
            if with_scores > 0:
                skills[skill] = skills[skill] / with_scores
        
        return jsonify({
            'success': True,
            'stats': {
                'total_interviews': total,
                'completed_interviews': completed,
                'average_score': round(avg_score, 1),
                'pass_rate': round((passed / with_scores * 100), 1) if with_scores > 0 else 0,
                'pending_analysis': completed - with_scores,
                'skills_average': skills
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting interview stats: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        session.close()


@app.route('/api/interview-automation/toggle', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=5, time_window=60)
def toggle_automation():
    """Start or stop the interview automation system"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        action = data.get('action', 'toggle')
        
        from interview_automation import interview_automation
        
        if action == 'start':
            start_interview_automation()
            message = "Interview automation started"
        elif action == 'stop':
            stop_interview_automation()
            message = "Interview automation stopped"
        else:
            # Toggle
            if interview_automation.is_running:
                stop_interview_automation()
                message = "Interview automation stopped"
            else:
                start_interview_automation()
                message = "Interview automation started"
        
        return jsonify({
            'success': True,
            'message': message,
            'is_running': interview_automation.is_running
        }), 200
        
    except Exception as e:
        logger.error(f"Error toggling automation: {e}")
        return jsonify({"error": str(e)}), 500

# 3. Fix the api_schedule_interview function to properly handle company_name
@app.route('/api/schedule-interview', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=10, time_window=60)
def api_schedule_interview():
    """Schedule interview with enhanced knowledge base creation for proactive questioning"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        data = request.json
        candidate_id = data.get('candidate_id')
        email = data.get('email')
        interview_date = data.get('date')
        time_slot = data.get('time_slot')
        job_description_override = data.get('job_description')
        
        logger.info(f"Schedule interview request: candidate_id={candidate_id}")
        
        if not candidate_id and not email:
            return jsonify({"success": False, "message": "candidate_id or email is required"}), 400
        
        session = SessionLocal()
        try:
            # Find candidate
            if candidate_id:
                candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            else:
                candidate = session.query(Candidate).filter_by(email=email).first()
            
            if not candidate:
                return jsonify({"success": False, "message": "Candidate not found"}), 404
            
            # Check if already scheduled
            if candidate.interview_scheduled and candidate.interview_token:
                existing_link = f"{request.host_url.rstrip('/')}/secure-interview/{candidate.interview_token}"
                return jsonify({
                    "success": True,
                    "message": "Interview already scheduled",
                    "interview_link": existing_link,
                    "knowledge_base_id": getattr(candidate, 'knowledge_base_id', None),
                    "already_scheduled": True
                }), 200
            
            # Extract resume content
            resume_content = ""
            resume_extracted = False
            
            if candidate.resume_path and os.path.exists(candidate.resume_path):
                logger.info(f"Extracting resume from: {candidate.resume_path}")
                resume_content = extract_resume_content(candidate.resume_path)
                if resume_content:
                    resume_extracted = True
                    logger.info(f"✅ Resume extracted: {len(resume_content)} characters")
                else:
                    logger.error("Resume extraction returned empty content")
            
            # Fallback to candidate profile if no resume
            if not resume_content:
                logger.warning("Using candidate profile as fallback")
                resume_content = f"""
CANDIDATE: {candidate.name}
EMAIL: {candidate.email}
POSITION: {candidate.job_title}
ATS SCORE: {candidate.ats_score}
STATUS: {candidate.status}
{f"SCORING: {candidate.score_reasoning}" if candidate.score_reasoning else ""}
"""
            
            # Get company name
            company_name = os.getenv('COMPANY_NAME', 'Our Company')
            
            # Get job description
            job_description = job_description_override or getattr(candidate, 'job_description', f"Position: {candidate.job_title}")
            
            # CREATE HEYGEN KNOWLEDGE BASE WITH INTERVIEW QUESTIONS
            knowledge_base_id = None
            kb_creation_method = "none"
            
            if os.getenv('HEYGEN_API_KEY') and resume_content:
                try:
                    logger.info("Creating HeyGen knowledge base with interview questions...")
                    
                    # Generate structured interview questions
                    interview_questions = generate_interview_questions(
                        candidate_name=candidate.name,
                        position=candidate.job_title,
                        resume_content=resume_content,
                        job_description=job_description
                    )
                    
                    kb_name = f"Interview_{candidate.name.replace(' ', '_')}_{candidate.id}"
                    
                    # Create comprehensive knowledge base content
                    kb_content = f"""
INTERVIEW CONFIGURATION:
- Mode: Structured Technical Interview
- Candidate: {candidate.name}
- Position: {candidate.job_title}
- Company: {company_name}
- Interview Type: Technical and Behavioral
- Duration: 30-45 minutes

SPECIAL COMMANDS:
- When you receive "INIT_INTERVIEW": Start with the warm greeting and first question
- When you receive "NEXT_QUESTION": Move to the next question in the list
- If user is silent for 15+ seconds: Gently prompt or ask if they need more time


CANDIDATE BACKGROUND:
{resume_content[:8000]}

JOB REQUIREMENTS:
{job_description[:2000]}

{interview_questions}

INTERVIEW BEHAVIOR INSTRUCTIONS:
1. When stream starts, wait for "INIT_INTERVIEW" command
2. Upon receiving "INIT_INTERVIEW", immediately greet the candidate and ask the first question
3. Listen to complete answers before proceeding
4. Ask follow-up questions when appropriate
5. Keep track of which questions you've asked
6. Be encouraging if candidate seems nervous
7. End professionally after covering all questions

CONVERSATION STARTERS:
- If you receive any greeting like "Hello", "Hi", respond with: "Hello {candidate.name}! Welcome to your interview for {candidate.job_title} at {company_name}. I'm excited to learn about your experience. Let's start with you telling me about yourself and your journey to applying for this role."
- If candidate asks "Can you hear me?", respond: "Yes, I can hear you clearly! Let's begin with our interview. Please tell me about yourself."
- If candidate seems confused, say: "No worries! This is an AI-powered interview. I'll be asking you questions about your experience and the {candidate.job_title} role. Shall we start?"

IMPORTANT RULES:
- Start immediately when you receive "INIT_INTERVIEW"
- Always maintain a professional yet friendly tone
- Give candidates time to think (10-15 seconds)
- If no response after 20 seconds, ask: "Take your time, or would you like me to rephrase the question?"
- Track answered questions to avoid repetition
"""                    
                    # Prepare HeyGen payload with proper configuration
                    heygen_payload = {
                        'name': kb_name,
                        'description': f'Structured interview for {candidate.name} - {candidate.job_title}',
                        'content': kb_content,
                        'opening_line': f"Hello {candidate.name}, welcome to your interview for the {candidate.job_title} position at {company_name}. I'm your AI interviewer today. I've reviewed your resume and I'm excited to learn more about your experiences. Let's start with you telling me a bit about yourself and your journey to applying for this role.",
                        'custom_prompt': f"""You are conducting a professional technical interview for {candidate.name}. 
                        
Your personality: Professional, friendly, encouraging, and engaged.

Key behaviors:
1. Ask questions from the provided list ONE AT A TIME
2. Wait for complete answers before proceeding
3. Show active listening with phrases like "That's interesting", "I see", "Tell me more"
4. If they struggle, offer encouragement: "Take your time", "No worries"
5. Ask follow-up questions based on their responses
6. Keep track of which questions you've asked to avoid repetition

Interview style:
- Conversational, not robotic
- Professional but warm
- Encouraging when candidate seems nervous
- Patient with responses

Remember: This is a conversation, not an interrogation. Make {candidate.name} feel comfortable while thoroughly assessing their qualifications for the {candidate.job_title} role."""
                    }
                    
                    # Make API call to HeyGen
                    heygen_response = requests.post(
                        'https://api.heygen.com/v1/streaming/knowledge_base',
                        headers={
                            'X-Api-Key': os.getenv('HEYGEN_API_KEY'),
                            'Content-Type': 'application/json',
                            'Accept': 'application/json'
                        },
                        json=heygen_payload,
                        timeout=30
                    )
                    
                    if heygen_response.ok:
                        kb_data = heygen_response.json()
                        knowledge_base_id = kb_data.get('data', {}).get('knowledge_base_id')
                        kb_creation_method = "heygen_api"
                        logger.info(f"✅ HeyGen KB created successfully: {knowledge_base_id}")
                    else:
                        error_text = heygen_response.text
                        logger.error(f"HeyGen API error: {heygen_response.status_code} - {error_text}")
                        
                except Exception as e:
                    logger.error(f"HeyGen KB creation failed: {e}", exc_info=True)
            
            # Fallback KB ID if HeyGen fails
            if not knowledge_base_id:
                knowledge_base_id = f"kb_{candidate.id}_{int(time.time())}"
                kb_creation_method = "fallback"
                logger.warning(f"Using fallback KB: {knowledge_base_id}")
            
            # Create interview session
            interview_token = str(uuid.uuid4())
            interview_session_id = f"session_{candidate.id}_{int(time.time())}"
            
            # Parse interview date
            if isinstance(interview_date, str):
                interview_datetime = datetime.fromisoformat(interview_date.replace('Z', '+00:00'))
            else:
                interview_datetime = datetime.now() + timedelta(days=3)
            
            # Update candidate record
            candidate.interview_scheduled = True
            candidate.interview_date = interview_datetime
            candidate.interview_token = interview_token
            candidate.interview_link = f"{request.host_url.rstrip('/')}/secure-interview/{interview_token}"
            candidate.final_status = 'Interview Scheduled'
            
            # Safe attribute setting for optional fields
            safe_attrs = {
                'interview_session_id': interview_session_id,
                'knowledge_base_id': knowledge_base_id,
                'interview_created_at': datetime.now(),
                'interview_expires_at': datetime.now() + timedelta(days=7),
                'company_name': company_name,
                'interview_time_slot': time_slot,
                'interview_questions_asked': '[]',
                'interview_answers_given': '[]',
                'interview_total_questions': 0,
                'interview_answered_questions': 0,
                'job_description': job_description if job_description_override else None
            }
            
            for attr, value in safe_attrs.items():
                if hasattr(candidate, attr):
                    setattr(candidate, attr, value)
            
            # Commit changes
            session.commit()
            
            # Send email
            email_sent = False
            try:
                send_interview_link_email(
                    candidate_email=candidate.email,
                    candidate_name=candidate.name,
                    interview_link=candidate.interview_link,
                    interview_date=interview_datetime,
                    time_slot=time_slot,
                    position=candidate.job_title
                )
                email_sent = True
                logger.info(f"Interview email sent to {candidate.email}")
            except Exception as e:
                logger.error(f"Email failed: {e}")
            
            # Clear caches
            cache.delete_memoized(get_cached_candidates)
            
            return jsonify({
                "success": True,
                "message": f"Interview scheduled for {candidate.name}",
                "interview_link": candidate.interview_link,
                "interview_date": interview_datetime.isoformat(),
                "knowledge_base_id": knowledge_base_id,
                "kb_creation_method": kb_creation_method,
                "resume_extracted": resume_extracted,
                "resume_content_length": len(resume_content),
                "email_sent": email_sent,
                "session_id": interview_session_id
            }), 200
            
        except Exception as e:
            session.rollback()
            logger.error(f"Error in schedule_interview: {e}", exc_info=True)
            return jsonify({"success": False, "message": str(e)}), 500
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Critical error: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500


def generate_interview_questions(candidate_name, position, resume_content, job_description):
    """Generate structured interview questions based on resume and job"""
    
    # Extract key skills from resume
    skills = extract_skills_from_resume(resume_content)
    
    questions = f"""
INTERVIEW QUESTIONS:

1. INTRODUCTION (Ask first):
   - "Tell me about yourself and your journey to applying for this {position} role."
   - "What attracted you to our company and this position?"

2. TECHNICAL QUESTIONS (Based on resume):"""
    
    # Add technical questions based on skills found
    if 'python' in resume_content.lower():
        questions += """
   - "I see you have Python experience. Can you tell me about a complex Python project you've worked on?"
   - "How do you handle error handling and debugging in Python?"""
   
    if 'javascript' in resume_content.lower() or 'react' in resume_content.lower():
        questions += """
   - "Tell me about your experience with JavaScript/React. What was the most challenging frontend problem you've solved?"
   - "How do you manage state in React applications?"""
    
    if 'database' in resume_content.lower() or 'sql' in resume_content.lower():
        questions += """
   - "Describe your experience with databases. How do you optimize slow queries?"
   - "Tell me about a time you designed a database schema."""
    
    questions += f"""

3. BEHAVIORAL QUESTIONS:
   - "Describe a time when you had to work under pressure. How did you handle it?"
   - "Tell me about a project where you had to collaborate with a difficult team member."
   - "Give me an example of when you had to learn a new technology quickly."

4. ROLE-SPECIFIC QUESTIONS:
   - "How do you see yourself contributing to our team in the first 90 days?"
   - "What aspects of this {position} role excite you the most?"

5. CLOSING QUESTIONS:
   - "What questions do you have for me about the role or the company?"
   - "Is there anything else you'd like me to know about your qualifications?"

REMEMBER: Ask these questions one at a time, wait for complete responses, and ask relevant follow-up questions based on their answers.
"""
    
    return questions


# Add this helper endpoint to check interview status
@app.route('/api/interview-status/<int:candidate_id>', methods=['GET'])
def get_interview_status(candidate_id):
    """Check if interview is already scheduled for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "exam_completed": candidate.exam_completed,
            "exam_percentage": candidate.exam_percentage,
            "interview_scheduled": candidate.interview_scheduled,
            "interview_token": candidate.interview_token,
            "interview_link": candidate.interview_link,
            "final_status": candidate.final_status
        }), 200
    finally:
        session.close()


@app.route('/api/verify-interview-process/<int:candidate_id>', methods=['GET'])
def verify_interview_process(candidate_id):
    """Verify the interview link generation process for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Check assessment completion
        assessment_status = {
            "exam_completed": candidate.exam_completed,
            "exam_percentage": candidate.exam_percentage,
            "exam_score": candidate.exam_score,
            "exam_feedback": candidate.exam_feedback
        }
        
        # Check interview status
        interview_status = {
            "interview_scheduled": candidate.interview_scheduled,
            "interview_token": candidate.interview_token,
            "interview_link": candidate.interview_link,
            "knowledge_base_id": candidate.knowledge_base_id,
            "final_status": candidate.final_status
        }
        
        # Determine if interview should be scheduled
        should_schedule = (
            candidate.exam_completed and 
            candidate.exam_percentage >= 70 and 
            not candidate.interview_scheduled
        )
        
        # Generate test interview link if needed
        test_link = None
        if should_schedule and not candidate.interview_token:
            import uuid
            test_token = str(uuid.uuid4())
            test_link = f"{request.host_url.rstrip('/')}/secure-interview/{test_token}"
        
        return jsonify({
            "candidate_info": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "job_title": candidate.job_title
            },
            "assessment_status": assessment_status,
            "interview_status": interview_status,
            "should_schedule_interview": should_schedule,
            "test_interview_link": test_link,
            "process_ready": candidate.exam_completed and candidate.exam_percentage is not None
        }), 200
    finally:
        session.close()

# Add these endpoints to your backend.py

# Add these endpoints to your backend.py

@app.route('/api/interview/recording/start', methods=['POST', 'OPTIONS'])
def start_recording():
    if request.method == 'OPTIONS':
        return _ok_preflight()

    body = request.get_json(silent=True) or {}
    session_id = body.get('session_id')
    if not session_id:
        return jsonify({"success": False, "error": "session_id required"}), 400

    base = _ensure_dir(os.path.join('logs', 'interviews', session_id))
    _append_jsonl(os.path.join(base, 'session.jsonl'), {
        "event": "recording_started",
        "ts": datetime.now(timezone.utc).isoformat(),
        "meta": {
            "recording_format": body.get('recording_format'),
            "client_meta": body.get('client_meta'),
        }
    })

    # optional: flag on candidate
    try:
        db = SessionLocal()
        try:
            cand = db.query(Candidate).filter_by(interview_session_id=session_id).first()
            if cand:
                cand.recording_started_at = datetime.now(timezone.utc)
                db.commit()
        finally:
            db.close()
    except Exception:
        logger.exception("start_recording: candidate update failed")

    return jsonify({"success": True}), 200

@app.route("/api/interview/recording/upload", methods=["POST", "OPTIONS"])
def upload_recording():
    # CORS preflight
    if request.method == "OPTIONS":
        return _ok_preflight()

    # multipart/form-data expected
    f = request.files.get("recording")
    session_id = request.form.get("session_id") or request.args.get("session_id")

    if not f or not session_id:
        return jsonify({"success": False, "error": "missing recording file or session_id"}), 400

    # sanitize session_id for filesystem usage
    safe_session = secure_filename(str(session_id)) or f"session_{int(time.time())}"

    # Decide filename + ext
    ext = _ext_from_filename(getattr(f, "filename", None), default_ext="webm")
    base = _ensure_dir(os.path.join("logs", "interviews", safe_session))
    fname = f"interview_{safe_session}_{int(time.time())}.{ext}"
    path = os.path.join(base, fname)

    try:
        # Save file
        f.save(path)

        # Log event
        _append_jsonl(os.path.join(base, "session.jsonl"), {
            "event": "recording_uploaded",
            "ts": datetime.now(timezone.utc).isoformat(),
            "file": path,
            "size": os.path.getsize(path),
            "session_id": safe_session,
        })

        # Optional: update candidate record
        try:
            db = SessionLocal()
            try:
                cand = db.query(Candidate).filter_by(interview_session_id=session_id).first()
                if cand:
                    cand.recording_path = path
                    cand.interview_recording_format = ext
                    cand.interview_recording_status = "completed"
                    # only set completed_at if not already present
                    if not getattr(cand, "interview_completed_at", None):
                        cand.interview_completed_at = datetime.now(timezone.utc)
                    db.commit()
            except Exception as e:
                db.rollback()
                logger.exception(f"Failed to update candidate recording info: {e}")
            finally:
                db.close()
        except Exception as e:
            logger.exception(f"upload_recording: candidate update failed: {e}")

        return jsonify({"success": True, "path": path}), 200

    except Exception as e:
        logger.exception(f"upload_recording: saving failed: {e}")
        return jsonify({"success": False, "error": "failed to save recording"}), 500

# @app.route('/api/interview/recording/upload', methods=['POST', 'OPTIONS'])
# def upload_recording():
    # if request.method == 'OPTIONS':
        # return _ok_preflight()
# 
    # multipart/form-data expected
    # f = request.files.get('recording')
    # session_id = request.form.get('session_id') or request.args.get('session_id')
    # if not f or not session_id:
        # return jsonify({"success": False, "error": "missing recording file or session_id"}), 400
# 
    # save file
    # base = _ensure_dir(os.path.join('logs', 'interviews', session_id))
    # fname = f"interview_{session_id}_{int(time.time())}.webm"
    # path = os.path.join(base, fname)
    # f.save(path)
# 
    # log event
    # _append_jsonl(os.path.join(base, 'session.jsonl'), {
        # "event": "recording_uploaded",
        # "ts": datetime.now(timezone.utc).isoformat(),
        # "file": path,
        # "size": os.path.getsize(path)
    # })
# 
    # optional: mark candidate
    # try:
        # db = SessionLocal()
        # try:
            # cand = db.query(Candidate).filter_by(interview_session_id=session_id).first()
            # if cand:
                # cand.recording_path = path
                # cand.interview_recording_format = 'webm'
                # cand.interview_recording_status = 'completed'
                # cand.interview_completed_at = cand.interview_completed_at or datetime.now(timezone.utc)
                # db.commit()
        # except Exception as e:
            # db.rollback()
            # logger.exception(f"Failed to update candidate recording info: {e}")
        # finally:
            # db.close()
    # except Exception as e:
        # logger.exception(f"upload_recording: candidate update failed: {e}")
# 
    # return jsonify({"success": True, "path": path}), 200

 
@app.route('/api/interview/qa/track', methods=['POST', 'OPTIONS'])
def track_qa_enhanced():
    """Enhanced Q&A tracking with better session management and error handling"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        interview_token = data.get('interview_token')
        candidate_id = data.get('candidate_id')
        qa_type = data.get('type')  # 'question', 'answer', or 'system'
        content = data.get('content')
        timestamp = data.get('timestamp')
        metadata = data.get('metadata', {})
        
        # Log the incoming request
        logger.info(f"Q&A Track Request: type={qa_type}, session={session_id}, token={interview_token}")
        
        # Validate required fields
        if not (session_id or interview_token) or not qa_type or not content:
            return jsonify({
                "error": "Missing required fields",
                "required": ["session_id OR interview_token", "type", "content"]
            }), 400
        
        session = SessionLocal()
        try:
            # Find candidate using multiple strategies
            candidate = None
            
            # Strategy 1: Try session ID
            if session_id and session_id != 'null' and session_id != 'undefined':
                candidate = session.query(Candidate).filter_by(
                    interview_session_id=session_id
                ).first()
                logger.info(f"Session ID lookup: {'found' if candidate else 'not found'}")
            
            # Strategy 2: Try interview token
            if not candidate and interview_token:
                candidate = session.query(Candidate).filter_by(
                    interview_token=interview_token
                ).first()
                logger.info(f"Token lookup: {'found' if candidate else 'not found'}")
            
            # Strategy 3: Try candidate ID as last resort
            if not candidate and candidate_id:
                try:
                    cid = int(str(candidate_id).replace('emergency_', '').split('_')[0])
                    candidate = session.query(Candidate).filter_by(id=cid).first()
                    logger.info(f"Candidate ID lookup: {'found' if candidate else 'not found'}")
                except:
                    pass
            
            if not candidate:
                logger.error(f"No candidate found for session={session_id}, token={interview_token}, id={candidate_id}")
                
                # If it's an emergency session, try to create a placeholder
                if session_id and 'emergency' in str(session_id):
                    return jsonify({
                        "error": "Emergency session - candidate not found",
                        "session_id": session_id,
                        "suggestion": "Refresh the page to reinitialize the interview"
                    }), 404
                
                return jsonify({"error": "Session not found"}), 404
            
            # Initialize Q&A storage if needed
            if not hasattr(candidate, 'interview_questions_asked'):
                candidate.interview_questions_asked = '[]'
            if not hasattr(candidate, 'interview_answers_given'):
                candidate.interview_answers_given = '[]'
            if not hasattr(candidate, 'interview_transcript'):
                candidate.interview_transcript = ''
            
            # Parse existing Q&A data
            try:
                questions = json.loads(candidate.interview_questions_asked or '[]')
            except:
                questions = []
                
            try:
                answers = json.loads(candidate.interview_answers_given or '[]')
            except:
                answers = []
            
            # Track based on type
            if qa_type == 'question':
                question_data = {
                    'id': str(uuid.uuid4()),
                    'text': content,
                    'timestamp': timestamp or datetime.now().isoformat(),
                    'order': len(questions) + 1,
                    'type': 'avatar_question',
                    'metadata': metadata
                }
                questions.append(question_data)
                candidate.interview_questions_asked = json.dumps(questions, ensure_ascii=False)
                candidate.interview_total_questions = len(questions)
                
                logger.info(f"Added question #{len(questions)}: {content[:50]}...")
                
            elif qa_type == 'answer':
                # Find which question this answers
                question_order = len(questions)  # Default to last question
                
                # Try to match with unanswered questions
                unanswered_count = len(questions) - len(answers)
                if unanswered_count > 0:
                    question_order = len(answers) + 1
                
                answer_data = {
                    'id': str(uuid.uuid4()),
                    'text': content,
                    'timestamp': timestamp or datetime.now().isoformat(),
                    'question_order': question_order,
                    'order': len(answers) + 1,
                    'type': 'candidate_answer',
                    'metadata': metadata
                }
                answers.append(answer_data)
                candidate.interview_answers_given = json.dumps(answers, ensure_ascii=False)
                candidate.interview_answered_questions = len(answers)
                
                logger.info(f"Added answer #{len(answers)}: {content[:50]}...")
                
            elif qa_type == 'system':
                # System messages (like initialization)
                logger.info(f"System message: {content}")
            
            # Update transcript
            timestamp_str = datetime.now().strftime('%H:%M:%S')
            role = 'AI' if qa_type == 'question' else 'Candidate' if qa_type == 'answer' else 'System'
            transcript_entry = f"\n[{timestamp_str}] {role}: {content}"
            
            if candidate.interview_transcript:
                candidate.interview_transcript += transcript_entry
            else:
                candidate.interview_transcript = transcript_entry
            
            # Update session tracking
            if session_id and (not candidate.interview_session_id or candidate.interview_session_id != session_id):
                candidate.interview_session_id = session_id
                logger.info(f"Updated session ID to: {session_id}")
            
            # Update activity timestamp
            candidate.last_accessed = datetime.now()
            
            # Commit changes
            session.commit()
            
            # Return comprehensive response
            return jsonify({
                "success": True,
                "candidate_id": candidate.id,
                "candidate_name": candidate.name,
                "session_id": candidate.interview_session_id,
                "total_questions": len(questions),
                "total_answers": len(answers),
                "unanswered_questions": len(questions) - len(answers),
                "message": f"{qa_type} tracked successfully",
                "debug": {
                    "found_by": "session_id" if session_id == candidate.interview_session_id else "token" if interview_token else "candidate_id",
                    "transcript_length": len(candidate.interview_transcript or ''),
                    "last_question": questions[-1]['text'][:50] if questions else None,
                    "last_answer": answers[-1]['text'][:50] if answers else None
                }
            }), 200
            
        except Exception as e:
            session.rollback()
            logger.error(f"Error in track_qa: {e}", exc_info=True)
            return jsonify({
                "error": f"Database error: {str(e)}",
                "type": type(e).__name__
            }), 500
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error tracking Q&A: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


# backend.py
@app.route('/api/resume-text/<int:candidate_id>', methods=['GET'])
def api_resume_text(candidate_id):
    session = SessionLocal()
    try:
        cand = session.query(Candidate).filter_by(id=candidate_id).first()
        if not cand:
            return jsonify({"error": "Candidate not found"}), 404

        text = ""
        if cand.resume_path and os.path.exists(cand.resume_path):
            # uses your existing extractor
            text = extract_resume_content(cand.resume_path)

        return jsonify({"resume_text": text, "length": len(text)}), 200
    finally:
        session.close()


@app.route('/api/interview/qa/verify/<session_id>', methods=['GET'])
def verify_qa_tracking_enhanced(session_id):
    """Enhanced verification of Q&A tracking with detailed analysis"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(
            interview_session_id=session_id
        ).first()
        
        if not candidate:
            # Try to find by token if session_id looks like it contains a token
            if 'token_' in session_id:
                token = session_id.split('token_')[-1]
                candidate = session.query(Candidate).filter_by(
                    interview_token=token
                ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Parse Q&A data
        try:
            questions = json.loads(candidate.interview_questions_asked or '[]')
        except:
            questions = []
            
        try:
            answers = json.loads(candidate.interview_answers_given or '[]')
        except:
            answers = []
        
        # Create detailed Q&A pairs analysis
        qa_pairs = []
        for i, question in enumerate(questions):
            # Find matching answer
            matching_answer = None
            for answer in answers:
                if answer.get('question_order') == i + 1:
                    matching_answer = answer
                    break
            
            # If no exact match, try to match by timing
            if not matching_answer and i < len(answers):
                matching_answer = answers[i]
            
            qa_pairs.append({
                'question_number': i + 1,
                'question': {
                    'text': question.get('text', ''),
                    'timestamp': question.get('timestamp', ''),
                    'id': question.get('id', '')
                },
                'answer': {
                    'text': matching_answer.get('text', '') if matching_answer else None,
                    'timestamp': matching_answer.get('timestamp', '') if matching_answer else None,
                    'id': matching_answer.get('id', '') if matching_answer else None
                } if matching_answer else None,
                'has_answer': matching_answer is not None,
                'time_to_answer': calculate_time_difference(
                    question.get('timestamp'),
                    matching_answer.get('timestamp')
                ) if matching_answer else None
            })
        
        # Calculate statistics
        total_questions = len(questions)
        total_answers = len(answers)
        answer_rate = (total_answers / total_questions * 100) if total_questions > 0 else 0
        
        # Get transcript preview
        transcript = candidate.interview_transcript or ''
        transcript_lines = transcript.strip().split('\n')[-10:]  # Last 10 lines
        
        return jsonify({
            "session_id": session_id,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "statistics": {
                "total_questions": total_questions,
                "total_answers": total_answers,
                "unanswered_questions": total_questions - total_answers,
                "answer_rate": f"{answer_rate:.1f}%",
                "transcript_length": len(transcript),
                "transcript_lines": len(transcript.strip().split('\n')) if transcript else 0
            },
            "qa_pairs": qa_pairs,
            "recent_transcript": transcript_lines,
            "session_info": {
                "interview_token": candidate.interview_token,
                "knowledge_base_id": candidate.knowledge_base_id,
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "last_accessed": candidate.last_accessed.isoformat() if hasattr(candidate, 'last_accessed') and candidate.last_accessed else None
            },
            "recording_status": getattr(candidate, 'interview_recording_status', 'unknown')
        }), 200
        
    except Exception as e:
        logger.error(f"Error in verify Q&A: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


def calculate_time_difference(timestamp1, timestamp2):
    """Calculate time difference between two timestamps in seconds"""
    if not timestamp1 or not timestamp2:
        return None
    try:
        t1 = datetime.fromisoformat(timestamp1.replace('Z', '+00:00'))
        t2 = datetime.fromisoformat(timestamp2.replace('Z', '+00:00'))
        return abs((t2 - t1).total_seconds())
    except:
        return None


# Add this endpoint to check all interview data for a candidate
@app.route('/api/interview/candidate/<int:candidate_id>/full-data', methods=['GET'])
def get_candidate_interview_data(candidate_id):
    """Get complete interview data for a candidate including Q&A"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Parse all interview data
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        
        return jsonify({
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "interview": {
                "scheduled": candidate.interview_scheduled,
                "token": candidate.interview_token,
                "session_id": candidate.interview_session_id,
                "knowledge_base_id": candidate.knowledge_base_id,
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None
            },
            "qa_data": {
                "questions": questions,
                "answers": answers,
                "total_questions": len(questions),
                "total_answers": len(answers),
                "completion_rate": f"{(len(answers) / len(questions) * 100) if questions else 0:.1f}%"
            },
            "transcript": {
                "content": candidate.interview_transcript,
                "length": len(candidate.interview_transcript or ''),
                "lines": len((candidate.interview_transcript or '').strip().split('\n'))
            },
            "recording": {
                "status": candidate.interview_recording_status,
                "file": candidate.interview_recording_file,
                "duration": candidate.interview_recording_duration
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting candidate interview data: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

def upload_to_cloud_storage(local_path, filename):
    """Upload recording to cloud storage (implement based on your provider)"""
    try:
        # Example for S3
        if os.getenv('AWS_ACCESS_KEY_ID'):
            import boto3
            s3 = boto3.client('s3')
            bucket = os.getenv('S3_BUCKET_NAME', 'interview-recordings')
            key = f"interviews/{filename}"
            
            s3.upload_file(local_path, bucket, key)
            return f"https://{bucket}.s3.amazonaws.com/{key}"
        
        # Example for Google Cloud Storage
        # if os.getenv('GOOGLE_APPLICATION_CREDENTIALS'):
        #     from google.cloud import storage
        #     client = storage.Client()
        #     bucket = client.bucket('interview-recordings')
        #     blob = bucket.blob(f"interviews/{filename}")
        #     blob.upload_from_filename(local_path)
        #     return blob.public_url
        
        return None
        
    except Exception as e:
        logger.error(f"Cloud upload failed: {e}")
        return None


# Add this route to check recording and Q&A data
@app.route('/api/interview/session/data/<session_id>', methods=['GET'])
def get_interview_session_data(session_id):
    """Get complete interview session data including recording and Q&A"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(
            interview_session_id=session_id
        ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Parse Q&A data
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        
        # Create Q&A pairs
        qa_pairs = []
        for i, question in enumerate(questions):
            answer = next((a for a in answers if a.get('question_order') == i + 1), None)
            qa_pairs.append({
                'question': question,
                'answer': answer
            })
        
        return jsonify({
            "session_id": session_id,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "recording": {
                "status": candidate.interview_recording_status,
                "file": candidate.interview_recording_file,
                "url": candidate.interview_recording_url,
                "duration": candidate.interview_recording_duration,
                "size": candidate.interview_recording_size,
                "format": candidate.interview_recording_format
            },
            "qa_data": {
                "total_questions": candidate.interview_total_questions,
                "total_answers": candidate.interview_answered_questions,
                "qa_pairs": qa_pairs,
                "raw_questions": questions,
                "raw_answers": answers
            },
            "ai_analysis": {
                "status": candidate.interview_ai_analysis_status,
                "overall_score": candidate.interview_ai_score,
                "technical_score": candidate.interview_ai_technical_score,
                "communication_score": candidate.interview_ai_communication_score,
                "feedback": candidate.interview_ai_overall_feedback
            },
            "timestamps": {
                "started": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None
            }
        }), 200
        
    finally:
        session.close()



# Interview Session Management Endpoints
# Add these endpoints to your backend.py file

def _ok_preflight():
    resp = jsonify({})
    resp.status_code = 200
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    return resp

def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path

def _append_jsonl(path, obj):
    with open(path, 'a', encoding='utf-8') as f:
        f.write(json.dumps(obj, ensure_ascii=False) + '\n')

# POST /api/interview/session/start
@app.route('/api/interview/session/start', methods=['POST', 'OPTIONS'])
def start_interview_session():
    if request.method == 'OPTIONS':
        return _ok_preflight()

    body = request.get_json(silent=True) or {}
    interview_token   = body.get('interview_token')
    candidate_id      = body.get('candidate_id')
    recording_config  = body.get('recording_config') or {}

    # make a session id if none provided
    session_id = body.get('session_id') or f"sess_{uuid.uuid4().hex[:8]}_{int(time.time())}"

    # persist basic linkage if we can find the candidate
    try:
        session = SessionLocal()
        try:
            candidate = None
            if candidate_id:
                candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            elif interview_token:
                candidate = session.query(Candidate).filter_by(interview_token=interview_token).first()

            if candidate:
                candidate.interview_session_id = session_id
                candidate.interview_started_at = datetime.now(timezone.utc)
                candidate.last_accessed = datetime.now(timezone.utc)
                # initialize counters if fields exist
                if not getattr(candidate, 'interview_total_questions', None):
                    candidate.interview_total_questions = 0
                if not getattr(candidate, 'interview_answered_questions', None):
                    candidate.interview_answered_questions = 0
                session.commit()
        finally:
            session.close()
    except Exception as e:
        logger.exception("start_interview_session: DB linkage failed")

    # create a place to log things to disk (optional but handy)
    base = _ensure_dir(os.path.join('logs', 'interviews', session_id))
    _append_jsonl(os.path.join(base, 'session.jsonl'), {
        "event": "session_started",
        "ts": datetime.now(timezone.utc).isoformat(),
        "session_id": session_id,
        "interview_token": interview_token,
        "candidate_id": candidate_id,
        "recording_config": recording_config,
    })

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Interview session initialized"
    }), 200


# 1. Add route to get candidate data by token
@app.route('/api/get-candidate-by-token/<token>', methods=['GET', 'OPTIONS'])
def get_candidate_by_token(token):
    """Get candidate data by interview token for KB creation"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        candidate_data = {
            "id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "job_title": candidate.job_title,
            "company": getattr(candidate, 'company_name', os.getenv('COMPANY_NAME', 'Our Company')),
            "resume_path": candidate.resume_path,
            "job_description": getattr(candidate, 'job_description', None),
            "ats_score": candidate.ats_score,
            "phone": getattr(candidate, 'phone', None)
        }
        
        return jsonify(candidate_data), 200
        
    except Exception as e:
        logger.error(f"Error getting candidate by token: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


# 2. Add route to extract resume content
@app.route('/api/extract-resume/<int:candidate_id>', methods=['GET', 'OPTIONS'])
def extract_resume_content_api(candidate_id):
    """Extract resume content for a candidate"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        if not candidate.resume_path or not os.path.exists(candidate.resume_path):
            return jsonify({
                "content": "",
                "message": "No resume file available",
                "file_path": candidate.resume_path
            }), 200
        
        # Extract resume content
        resume_content = extract_resume_content(candidate.resume_path)
        
        # Extract additional metadata
        skills = extract_skills_from_resume(resume_content)
        experience = extract_experience_years(resume_content)
        projects = extract_projects_from_resume(resume_content)
        
        return jsonify({
            "content": resume_content,
            "skills": skills,
            "experience": experience,
            "projects": projects,
            "file_path": candidate.resume_path,
            "content_length": len(resume_content),
            "extraction_successful": len(resume_content) > 0
        }), 200
        
    except Exception as e:
        logger.error(f"Error extracting resume: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

# 3. Enhanced knowledge base storage
@app.route('/api/store-knowledge-base', methods=['POST', 'OPTIONS'])
def store_knowledge_base_enhanced():
    """Store knowledge base data with enhanced tracking"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        candidate_id = data.get('candidate_id')
        kb_id = data.get('knowledge_base_id')
        content = data.get('content', '')
        
        if not candidate_id or not kb_id:
            return jsonify({"error": "candidate_id and knowledge_base_id are required"}), 400
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404
            
            # Update candidate with knowledge base information
            candidate.knowledge_base_id = kb_id
            candidate.interview_kb_id = kb_id  # Also store in interview-specific field
            
            # Store knowledge base content if we have the field
            if hasattr(candidate, 'interview_kb_content'):
                candidate.interview_kb_content = content
            
            # Store metadata
            if hasattr(candidate, 'interview_kb_metadata'):
                metadata = {
                    'created_at': data.get('created_at', datetime.now().isoformat()),
                    'content_length': len(content),
                    'kb_type': 'heygen' if not kb_id.startswith('kb_') else 'fallback',
                    'version': '1.0'
                }
                candidate.interview_kb_metadata = json.dumps(metadata)
            
            # Update timestamps
            candidate.interview_created_at = datetime.now()
            if not candidate.interview_expires_at:
                candidate.interview_expires_at = datetime.now() + timedelta(days=7)
            
            session.commit()
            
            logger.info(f"Knowledge base {kb_id} stored for candidate {candidate.name}")
            
            return jsonify({
                "success": True,
                "message": "Knowledge base stored successfully",
                "candidate_id": candidate_id,
                "knowledge_base_id": kb_id
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error storing knowledge base: {e}")
        return jsonify({"error": str(e)}), 500


# 4. Enhanced resume extraction function
def extract_resume_content(resume_path):
    """Enhanced resume extraction with fixed f-string syntax"""
    try:
        if not os.path.exists(resume_path):
            logger.error(f"Resume file not found: {resume_path}")
            return ""
        
        file_ext = os.path.splitext(resume_path)[1].lower()
        logger.info(f"Extracting resume: {resume_path} (type: {file_ext})")
        
        resume_text = ""
        
        if file_ext == '.pdf':
            # Try multiple PDF extraction methods
            resume_text = extract_pdf_content(resume_path)
        elif file_ext in ['.docx', '.doc']:
            resume_text = extract_docx_content(resume_path)
        elif file_ext == '.txt':
            resume_text = extract_txt_content(resume_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
        
        if resume_text and len(resume_text.strip()) > 50:
            logger.info(f"Resume extracted successfully: {len(resume_text)} characters")
            return resume_text.strip()
        else:
            logger.error("Resume extraction failed or produced minimal content")
            return ""
            
    except Exception as e:
        logger.error(f"Resume extraction error: {str(e)}", exc_info=True)
        return ""


def extract_pdf_content(pdf_path):
    """Extract content from PDF using multiple methods"""
    try:
        # Method 1: Try PyPDF2
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                if text.strip():
                    return text.strip()
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        # Method 2: Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # Method 3: Try pymupdf (fitz)
        try:
            import fitz  # pymupdf
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning(f"pymupdf extraction failed: {str(e)}")
            
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
    
    return ""


def extract_docx_content(docx_path):
    """Extract content from DOCX files"""
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text.strip()
        
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return ""


def extract_txt_content(txt_path):
    """Extract content from text files"""
    try:
        # Try UTF-8 first
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Fallback to latin-1
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
    except Exception as e:
        logger.error(f"TXT extraction failed: {str(e)}")
        return ""

# 5. Add validation endpoint for knowledge base creation
@app.route('/api/validate-kb-creation/<int:candidate_id>', methods=['GET', 'OPTIONS'])
def validate_kb_creation(candidate_id):
    """Validate that all components are ready for KB creation"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Check all requirements
        checks = {
            "candidate_exists": True,
            "interview_token_exists": bool(candidate.interview_token),
            "resume_file_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "heygen_api_configured": bool(os.getenv('HEYGEN_API_KEY')),
            "job_title_available": bool(candidate.job_title),
            "company_name_available": bool(getattr(candidate, 'company_name', None) or os.getenv('COMPANY_NAME'))
        }
        
        # Try to extract resume content
        resume_content = ""
        if checks["resume_file_exists"]:
            resume_content = extract_resume_content_enhanced(candidate.resume_path)
            checks["resume_extractable"] = len(resume_content) > 0
        else:
            checks["resume_extractable"] = False
        
        # Check if KB already exists
        checks["kb_already_exists"] = bool(candidate.knowledge_base_id)
        
        # Overall readiness
        critical_checks = ["candidate_exists", "interview_token_exists", "heygen_api_configured", "job_title_available"]
        all_critical_passed = all(checks[check] for check in critical_checks)
        
        return jsonify({
            "candidate_id": candidate_id,
            "candidate_name": candidate.name,
            "ready_for_kb_creation": all_critical_passed,
            "checks": checks,
            "resume_content_length": len(resume_content),
            "existing_kb_id": candidate.knowledge_base_id,
            "recommendations": generate_kb_recommendations(checks)
        }), 200
        
    finally:
        session.close()


def generate_kb_recommendations(checks):
    """Generate recommendations based on validation checks"""
    recommendations = []
    
    if not checks.get("heygen_api_configured"):
        recommendations.append("Set HEYGEN_API_KEY environment variable")
    
    if not checks.get("resume_file_exists"):
        recommendations.append("Upload candidate resume file")
    elif not checks.get("resume_extractable"):
        recommendations.append("Resume file exists but content extraction failed - check file format")
    
    if not checks.get("interview_token_exists"):
        recommendations.append("Generate interview token for candidate")
    
    if not checks.get("company_name_available"):
        recommendations.append("Set COMPANY_NAME environment variable or add company to candidate record")
    
    if checks.get("kb_already_exists"):
        recommendations.append("Knowledge base already exists - consider updating instead of recreating")
    
    return recommendations


# 6. Add test endpoint for KB creation
@app.route('/api/test-kb-creation/<int:candidate_id>', methods=['POST', 'OPTIONS'])
def test_kb_creation(candidate_id):
    """Test knowledge base creation for debugging"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Step 1: Validate candidate
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404
            
            test_results = {
                "candidate_info": {
                    "id": candidate.id,
                    "name": candidate.name,
                    "email": candidate.email,
                    "job_title": candidate.job_title
                }
            }
            
            # Step 2: Test resume extraction
            resume_content = ""
            if candidate.resume_path:
                resume_content = extract_resume_content_enhanced(candidate.resume_path)
                test_results["resume_extraction"] = {
                    "success": len(resume_content) > 0,
                    "content_length": len(resume_content),
                    "file_path": candidate.resume_path,
                    "preview": resume_content[:200] + "..." if resume_content else "No content"
                }
            
            # Step 3: Test HeyGen API connection
            heygen_key = os.getenv('HEYGEN_API_KEY')
            if heygen_key:
                try:
                    # Test with a simple knowledge base
                    test_kb_content = f"Test knowledge base for {candidate.name}"
                    
                    test_response = requests.post(
                        'https://api.heygen.com/v1/streaming/knowledge_base',
                        headers={
                            'X-Api-Key': heygen_key,
                            'Content-Type': 'application/json'
                        },
                        json={
                            'name': f'Test_KB_{candidate.id}_{int(time.time())}',
                            'description': 'Test knowledge base',
                            'content': test_kb_content
                        },
                        timeout=30
                    )
                    
                    test_results["heygen_api_test"] = {
                        "success": test_response.ok,
                        "status_code": test_response.status_code,
                        "response_preview": test_response.text[:500]
                    }
                    
                    if test_response.ok:
                        kb_data = test_response.json()
                        test_kb_id = kb_data.get('data', {}).get('knowledge_base_id')
                        if test_kb_id:
                            test_results["heygen_api_test"]["kb_id_created"] = test_kb_id
                
                except Exception as e:
                    test_results["heygen_api_test"] = {
                        "success": False,
                        "error": str(e)
                    }
            else:
                test_results["heygen_api_test"] = {
                    "success": False,
                    "error": "HEYGEN_API_KEY not configured"
                }
            
            return jsonify({
                "success": True,
                "test_results": test_results,
                "ready_for_production": (
                    test_results.get("heygen_api_test", {}).get("success", False) and
                    test_results.get("resume_extraction", {}).get("success", False)
                )
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"KB creation test failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/interview/question/add', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=50, time_window=60)
def add_interview_question():
    """Add a question asked by the avatar during interview"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        question_data = data.get('question_data', {})
        
        if not session_id or not question_data.get('text'):
            return jsonify({"success": False, "message": "session_id and question_data.text are required"}), 400
        
        # Import session manager
        from interview_session_manager import interview_session_manager
        
        # Add question to session
        question_id = interview_session_manager.add_interview_question(session_id, question_data)
        
        if question_id:
            return jsonify({
                "success": True,
                "question_id": question_id,
                "message": "Question added successfully"
            }), 200
        else:
            return jsonify({"success": False, "message": "Failed to add question"}), 500
            
    except Exception as e:
        logger.error(f"Error adding interview question: {e}")
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/interview/answer/add', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=50, time_window=60)
def add_interview_answer():
    """Add a candidate's answer during interview"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        question_id = data.get('question_id')
        answer_data = data.get('answer_data', {})
        
        if not session_id or not question_id or not answer_data.get('text'):
            return jsonify({"success": False, "message": "session_id, question_id, and answer_data.text are required"}), 400
        
        # Import session manager
        from interview_session_manager import interview_session_manager
        
        # Add answer to session
        answer_id = interview_session_manager.add_interview_answer(session_id, question_id, answer_data)
        
        if answer_id:
            return jsonify({
                "success": True,
                "answer_id": answer_id,
                "message": "Answer added successfully"
            }), 200
        else:
            return jsonify({"success": False, "message": "Failed to add answer"}), 500
            
    except Exception as e:
        logger.error(f"Error adding interview answer: {e}")
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/interview/session/end', methods=['POST', 'OPTIONS'])
def end_interview_session():
    """End the interview session"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        
        if not session_id:
            return jsonify({"error": "session_id is required"}), 400
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(
                interview_session_id=session_id
            ).first()
            
            if not candidate:
                return jsonify({"error": "Session not found"}), 404
            
            # Update session status
            candidate.interview_completed_at = datetime.now()
            candidate.interview_status = 'completed'
            
            # Calculate duration
            if candidate.interview_started_at:
                duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                candidate.interview_duration = int(duration)
            
            session.commit()
            
            # Clear caches
            cache.delete_memoized(get_cached_candidates)
            
            logger.info(f"Interview session ended: {session_id}")
            
            return jsonify({
                "success": True,
                "message": "Interview session ended successfully",
                "session_id": session_id,
                "duration": getattr(candidate, 'interview_duration', 0)
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error ending interview session: {e}")
        return jsonify({"error": str(e)}), 500

# @app.route('/api/interview/session/<session_id>', methods=['GET'])
# def get_interview_session_data(session_id):
    # """Get complete interview session data"""
    # try:
        # Import session manager
        # from interview_session_manager import interview_session_manager
        # 
        # Get session data
        # session_data = interview_session_manager.get_session_data(session_id)
        # 
        # if session_data:
            # return jsonify({
                # "success": True,
                # "session_data": session_data
            # }), 200
        # else:
            # return jsonify({"success": False, "message": "Session not found"}), 404
            # 
    # except Exception as e:
        # logger.error(f"Error getting session data: {e}")
        # return jsonify({"success": False, "message": str(e)}), 500
# 

@app.route('/api/interview/analysis/<int:candidate_id>', methods=['GET'])
def get_interview_analysis(candidate_id):
    """Get AI analysis results for a candidate's interview"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        analysis_data = {
            "technical_score": candidate.interview_ai_technical_score,
            "communication_score": candidate.interview_ai_communication_score,
            "problem_solving_score": candidate.interview_ai_problem_solving_score,
            "cultural_fit_score": candidate.interview_ai_cultural_fit_score,
            "overall_score": candidate.interview_ai_score,
            "overall_feedback": candidate.interview_ai_overall_feedback,
            "question_analysis": json.loads(candidate.interview_ai_questions_analysis) if candidate.interview_ai_questions_analysis else [],
            "analysis_status": candidate.interview_ai_analysis_status,
            "final_status": candidate.interview_final_status
        }
        
        return jsonify({
            "success": True,
            "candidate_id": candidate_id,
            "analysis": analysis_data
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting interview analysis: {e}")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        session.close()

# Add this to your backend.py to properly create knowledge base from resume/job

@app.route('/api/create-interview-knowledge-base', methods=['POST', 'OPTIONS'])
def create_interview_knowledge_base():
    """Create HeyGen knowledge base from candidate's resume and job description"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        candidate_id = data.get('candidate_id')
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404
            
            # Extract resume content
            resume_content = ""
            if candidate.resume_path and os.path.exists(candidate.resume_path):
                resume_content = extract_resume_content(candidate.resume_path)
            
            # Get job description
            job_description = candidate.job_description or f"Position: {candidate.job_title}"
            
            # Create knowledge base content
            kb_content = f"""
            CANDIDATE INFORMATION:
            Name: {candidate.name}
            Email: {candidate.email}
            Position Applied: {candidate.job_title}
            
            RESUME CONTENT:
            {resume_content}
            
            JOB DESCRIPTION:
            {job_description}
            
            INTERVIEW INSTRUCTIONS:
            - Ask questions based on the candidate's experience mentioned in resume
            - Focus on skills required for {candidate.job_title}
            - Assess technical competence based on job requirements
            - Ask behavioral questions related to their past experiences
            """
            
            # Call HeyGen API to create knowledge base
            heygen_key = os.getenv('HEYGEN_API_KEY')
            if heygen_key:
                response = requests.post(
                    'https://api.heygen.com/v1/streaming/knowledge_base',
                    headers={
                        'X-Api-Key': heygen_key,
                        'Content-Type': 'application/json'
                    },
                    json={
                        'name': f"Interview_{candidate.name}_{candidate.job_title}",
                        'content': kb_content,
                        'custom_prompt': generate_custom_interview_prompt(candidate, resume_content, job_description)
                    }
                )
                
                if response.ok:
                    kb_data = response.json()
                    kb_id = kb_data['data']['knowledge_base_id']
                    
                    # Update candidate record
                    candidate.knowledge_base_id = kb_id
                    candidate.interview_kb_id = kb_id
                    session.commit()
                    
                    return jsonify({
                        "success": True,
                        "knowledge_base_id": kb_id
                    }), 200
            
            # Fallback if HeyGen unavailable
            fallback_kb_id = f"kb_{candidate_id}_{int(time.time())}"
            candidate.knowledge_base_id = fallback_kb_id
            session.commit()
            
            return jsonify({
                "success": True,
                "knowledge_base_id": fallback_kb_id,
                "fallback": True
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error creating knowledge base: {e}")
        return jsonify({"error": str(e)}), 500


def generate_custom_interview_prompt(candidate, resume_content, job_description):
    """Generate custom interview prompt based on resume and job"""
    return f"""
    You are interviewing {candidate.name} for {candidate.job_title} position.
    
    Based on their resume, ask questions about:
    1. Their experience with technologies mentioned in their resume
    2. Projects they've worked on
    3. Challenges they've faced
    4. Their approach to problem-solving
    
    Based on the job requirements, assess:
    1. Technical skills required for the role
    2. Soft skills and communication
    3. Cultural fit
    4. Career goals alignment
    
    Keep the interview conversational and professional.
    Ask follow-up questions based on their responses.
    """


# 2. Enhanced resume extraction function with better error handling
def extract_resume_content(resume_path):
    """Extract text content from resume with better error handling"""
    try:
        if not os.path.exists(resume_path):
            logger.error(f"Resume file not found: {resume_path}")
            return ""
        
        file_ext = os.path.splitext(resume_path)[1].lower()
        logger.info(f"Extracting resume: {resume_path} (type: {file_ext})")
        
        resume_text = ""
        
        if file_ext == '.pdf':
            # Try PyPDF2 first
            try:
                import PyPDF2
                with open(resume_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        resume_text += page.extract_text() + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with PyPDF2: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
            
            # Try pdfplumber as fallback
            try:
                import pdfplumber
                with pdfplumber.open(resume_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            resume_text += page_text + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with pdfplumber: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
                    
        elif file_ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(resume_path)
                resume_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                logger.info(f"DOCX extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                
        elif file_ext == '.txt':
            try:
                with open(resume_path, 'r', encoding='utf-8') as file:
                    resume_text = file.read()
                logger.info(f"TXT extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except UnicodeDecodeError:
                with open(resume_path, 'r', encoding='latin-1') as file:
                    resume_text = file.read()
                return resume_text.strip()
        
        # If we couldn't extract anything, return empty string
        if not resume_text:
            logger.error(f"Failed to extract any text from {resume_path}")
            
    except Exception as e:
        logger.error(f"Resume extraction failed: {e}", exc_info=True)
    
    return ""

# Add this fixed version to your backend.py

@app.route('/api/verify-kb-creation/<int:candidate_id>', methods=['GET'])
def verify_kb_creation(candidate_id):
    """Verify knowledge base creation for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Check if we can extract resume
        resume_content = ""
        if candidate.resume_path and os.path.exists(candidate.resume_path):
            resume_content = extract_resume_content(candidate.resume_path)
        
        # Safely get knowledge_base_id using getattr
        knowledge_base_id = getattr(candidate, 'knowledge_base_id', None)
        interview_scheduled = getattr(candidate, 'interview_scheduled', False)
        job_description_available = bool(getattr(candidate, 'job_description', None))
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "resume_path": candidate.resume_path,
            "resume_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "resume_content_length": len(resume_content),
            "resume_preview": resume_content[:500] + "..." if resume_content else "No content",
            "knowledge_base_id": knowledge_base_id,
            "interview_scheduled": interview_scheduled,
            "job_description_available": job_description_available,
            "database_ready": True
        }), 200
    except Exception as e:
        logger.error(f"Error in verify_kb_creation: {e}", exc_info=True)
        return jsonify({
            "error": str(e),
            "message": "Database schema may need updating. Run migration script."
        }), 500
    finally:
        session.close()

@app.route('/api/force-create-kb/<int:candidate_id>', methods=['POST'])
def force_create_kb(candidate_id):
    """Force create knowledge base for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Extract resume
        resume_content = ""
        if candidate.resume_path and os.path.exists(candidate.resume_path):
            resume_content = extract_resume_content(candidate.resume_path)
        
        # Try to create HeyGen KB
        kb_id = create_heygen_knowledge_base(
            candidate_name=candidate.name,
            position=candidate.job_title or "Software Engineer",
            resume_content=resume_content,
            company=os.getenv('COMPANY_NAME', 'Our Company')
        )
        
        # Check if HeyGen succeeded
        method = "heygen"
        if not kb_id or kb_id.startswith('kb_fallback'):
            method = "fallback"
            if not kb_id:
                kb_id = f"kb_fallback_{candidate.id}_{int(time.time())}"
        
        # Update candidate
        candidate.interview_kb_id = kb_id
        session.commit()
        
        # Clear cache
        cache.delete_memoized(get_cached_candidates)
        
        return jsonify({
            "success": True,
            "knowledge_base_id": kb_id,
            "resume_extracted": len(resume_content) > 0,
            "method": method
        }), 200
        
    finally:
        session.close()

# Add these endpoints to your backend.py file

@app.route('/api/fix-missing-knowledge-bases', methods=['POST'])
def fix_missing_knowledge_bases():
    """Fix all scheduled interviews that are missing knowledge bases"""
    session = SessionLocal()
    fixed_count = 0
    heygen_count = 0
    fallback_count = 0
    errors = []
    
    try:
        # Find all candidates with scheduled interviews but no KB
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True,
            Candidate.interview_kb_id.is_(None)
        ).all()
        
        logger.info(f"Found {len(candidates)} candidates with missing knowledge bases")
        
        for candidate in candidates:
            try:
                # Extract resume content
                resume_content = ""
                if candidate.resume_path and os.path.exists(candidate.resume_path):
                    resume_content = extract_resume_content(candidate.resume_path)
                    logger.info(f"Extracted {len(resume_content)} chars from resume for {candidate.name}")
                
                # Try to create HeyGen KB
                kb_id = create_heygen_knowledge_base(
                    candidate_name=candidate.name,
                    position=candidate.job_title,
                    resume_content=resume_content,
                    company=os.getenv('COMPANY_NAME', 'Our Company')
                )
                
                if kb_id and not kb_id.startswith('kb_fallback'):
                    heygen_count += 1
                    logger.info(f"✅ Created HeyGen KB for {candidate.name}: {kb_id}")
                else:
                    # Use fallback if HeyGen failed
                    kb_id = f"kb_fallback_{candidate.id}_{int(time.time())}"
                    fallback_count += 1
                    logger.warning(f"Using fallback KB for {candidate.name}")
                
                # Update candidate
                candidate.interview_kb_id = kb_id
                fixed_count += 1
                
            except Exception as e:
                error_msg = f"Failed to fix {candidate.name} (ID: {candidate.id}): {str(e)}"
                logger.error(error_msg, exc_info=True)
                errors.append(error_msg)
        
        # Commit all changes
        session.commit()
        
        # Clear cache
        cache.delete_memoized(get_cached_candidates)
        
        return jsonify({
            "success": True,
            "fixed_count": fixed_count,
            "heygen_created": heygen_count,
            "fallback_used": fallback_count,
            "total_found": len(candidates),
            "errors": errors,
            "message": f"Fixed {fixed_count} candidates ({heygen_count} HeyGen, {fallback_count} fallback)"
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Batch KB fix failed: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/debug/heygen-test-fixed', methods=['GET'])
def debug_heygen_test_fixed():
    """Test HeyGen API with corrected field names"""
    heygen_key = os.getenv('HEYGEN_API_KEY')
    
    if not heygen_key:
        return jsonify({"error": "HEYGEN_API_KEY not found"}), 400
    
    # Test payload with CORRECT field names
    test_payload = {
        'name': f'Test_KB_{int(time.time())}',
        'opening': 'Hello, this is a test interview. Please tell me about yourself.',  # FIXED
        'prompt': 'Test interview questions: 1. Tell me about yourself. 2. Why this role? 3. What are your strengths?'
    }
    
    try:
        response = requests.post(
            "https://api.heygen.com/v1/streaming/knowledge_base/create",
            headers={
                'X-Api-Key': heygen_key,
                'Content-Type': 'application/json'
            },
            json=test_payload,
            timeout=30
        )
        
        if response.ok:
            data = response.json()
            kb_id = (
                data.get('data', {}).get('knowledge_base_id') or
                data.get('data', {}).get('id') or
                data.get('knowledge_base_id') or
                data.get('id')
            )
            
            return jsonify({
                "success": True,
                "status_code": response.status_code,
                "response": data,
                "knowledge_base_id": kb_id,
                "message": "HeyGen KB created successfully!"
            }), 200
        else:
            return jsonify({
                "success": False,
                "status_code": response.status_code,
                "error": response.text
            }), 400
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/api/check-interview-issues', methods=['GET'])
def check_interview_issues():
    """Check for all interview-related issues"""
    session = SessionLocal()
    try:
        issues = {
            "missing_kb": [],
            "missing_token": [],
            "missing_resume": [],
            "expired_interviews": []
        }
        
        # Get all scheduled interviews
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        for candidate in candidates:
            candidate_info = {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "job_title": candidate.job_title
            }
            
            # Check for missing KB using interview_kb_id
            if not candidate.interview_kb_id:
                issues["missing_kb"].append(candidate_info)
            
            # Check for missing token
            if not candidate.interview_token:
                issues["missing_token"].append(candidate_info)
            
            # Check for missing resume
            if not candidate.resume_path or not os.path.exists(candidate.resume_path):
                issues["missing_resume"].append(candidate_info)
            
            # Check for expired interviews
            if candidate.interview_expires_at and candidate.interview_expires_at < datetime.now():
                issues["expired_interviews"].append(candidate_info)
        
        summary = {
            "total_scheduled_interviews": len(candidates),
            "issues_found": {
                "missing_knowledge_bases": len(issues["missing_kb"]),
                "missing_tokens": len(issues["missing_token"]),
                "missing_resumes": len(issues["missing_resume"]),
                "expired_interviews": len(issues["expired_interviews"])
            },
            "details": issues
        }
        
        return jsonify(summary), 200
        
    except Exception as e:
        logger.error(f"Error checking interview issues: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/debug/heygen-test', methods=['GET'])
def debug_heygen_test():
    """Test HeyGen API connection and knowledge base creation"""
    heygen_key = os.getenv('HEYGEN_API_KEY')
    
    # Check if API key exists
    if not heygen_key:
        return jsonify({
            "error": "HEYGEN_API_KEY not found in environment variables",
            "fix": "Add HEYGEN_API_KEY to your .env file"
        }), 400
    
    # Test payload
    test_payload = {
        'name': f'Test_KB_{int(time.time())}',
        'description': 'Test knowledge base creation',
        'content': 'Test interview questions: 1. Tell me about yourself. 2. Why this role?',
        'opening_line': 'Hello, this is a test interview.'
    }
    
    # Try different endpoints
    endpoints = [
        "https://api.heygen.com/v1/streaming/knowledge_base/create",
        "https://api.heygen.com/v1/streaming/knowledge_base",
        "https://api.heygen.com/v1/streaming_avatar/knowledge_base",
        "https://api.heygen.com/v1/knowledge_base"
    ]
    
    results = []
    
    for endpoint in endpoints:
        try:
            response = requests.post(
                endpoint,
                headers={
                    'X-Api-Key': heygen_key,
                    'Content-Type': 'application/json',
                    'Accept': 'application/json'
                },
                json=test_payload,
                timeout=30
            )
            
            result = {
                "endpoint": endpoint,
                "status_code": response.status_code,
                "success": response.ok,
                "response": response.text[:500] if not response.ok else response.json()
            }
            results.append(result)
            
            if response.ok:
                # Try to extract KB ID
                data = response.json()
                kb_id = (
                    data.get('data', {}).get('knowledge_base_id') or
                    data.get('knowledge_base_id') or
                    data.get('id') or
                    data.get('data', {}).get('id')
                )
                if kb_id:
                    result["knowledge_base_id"] = kb_id
                    
        except Exception as e:
            results.append({
                "endpoint": endpoint,
                "error": str(e)
            })
    
    return jsonify({
        "api_key_length": len(heygen_key),
        "api_key_preview": f"{heygen_key[:10]}...{heygen_key[-5:]}",
        "test_results": results
    }), 200

def create_heygen_knowledge_base(candidate_name, position, resume_content, company):
    """Create HeyGen knowledge base with correct field names"""
    
    heygen_key = os.getenv('HEYGEN_API_KEY')
    if not heygen_key:
        logger.error("HEYGEN_API_KEY not set!")
        return None
    
    # Extract skills for better questions
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    
    # Create a more HeyGen-friendly prompt format
    heygen_prompt = f"""You are an AI interviewer conducting a professional technical interview.

IMPORTANT: You must ask these exact questions in order when conducting the interview.

Candidate: {candidate_name}
Position: {position}
Company: {company}

When the interview starts or when you hear "INIT_INTERVIEW", immediately greet the candidate and ask Question 1.

INTERVIEW QUESTIONS TO ASK IN ORDER:

Question 1: "Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background?"

Question 2: {"I see from your resume that you have experience with " + skills[0] + ". Can you tell me about a specific project where you used this technology?" if skills else "Can you tell me about your most significant technical project and the technologies you used?"}

Question 3: "Can you describe a time when you encountered a complex technical problem and walk me through how you approached and solved it?"

Question 4: "Tell me about a time when you had to collaborate with team members on a challenging project. How did you handle any conflicts?"

Question 5: "What would you say is your strongest technical skill, and can you give me a detailed example of how you've applied it?"

Question 6: "Technology evolves rapidly. Can you tell me about a time when you had to quickly learn a new technology for a project?"

Question 7: "Based on your understanding of this {position} role, how do you see your skills contributing to our team?"

Question 8: "What do you think would be the biggest challenge for you in this role?"

Question 9: "Where do you see your career heading in the next 3-5 years?"

Question 10: "Do you have any questions for me about the role or {company}?"

INSTRUCTIONS:
- Ask ONE question at a time
- Wait for complete answers before proceeding
- Be professional and encouraging
- If the candidate seems stuck, offer to rephrase the question
- After the last question, thank them for their time"""
    
    # HeyGen payload with CORRECT field names
    payload = {
        "name": f"Interview_{candidate_name.replace(' ', '_')}_{int(time.time())}",
        "opening": f"Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background?",
        "prompt": heygen_prompt
    }
    
    try:
        logger.info(f"Creating KB for {candidate_name} with payload keys: {payload.keys()}")
        
        response = requests.post(
            "https://api.heygen.com/v1/streaming/knowledge_base/create",
            headers={
                "X-Api-Key": heygen_key,
                "Content-Type": "application/json"
            },
            json=payload,
            timeout=45
        )
        
        logger.info(f"HeyGen response: {response.status_code}")
        
        if response.ok:
            data = response.json()
            logger.info(f"HeyGen success response: {json.dumps(data, indent=2)}")
            
            # Extract KB ID from response
            kb_id = None
            if isinstance(data, dict):
                kb_id = (
                    data.get('data', {}).get('knowledge_base_id') or
                    data.get('data', {}).get('id') or
                    data.get('knowledge_base_id') or
                    data.get('id')
                )
            
            if kb_id:
                logger.info(f"✅ Successfully created HeyGen KB: {kb_id}")
                return kb_id
            else:
                logger.error(f"KB ID not found in response: {data}")
                
        else:
            error_text = response.text
            logger.error(f"HeyGen API error {response.status_code}: {error_text}")
            
    except Exception as e:
        logger.error(f"HeyGen API exception: {type(e).__name__}: {str(e)}")
    
    return None

# Add to backend.py

@app.route('/api/interview/realtime-analysis', methods=['GET'])
def get_realtime_analysis_status():
    """Get real-time analysis status for all pending interviews"""
    session = SessionLocal()
    try:
        pending = session.query(Candidate).filter(
            Candidate.interview_completed_at.isnot(None),
            Candidate.interview_ai_analysis_status.in_(['pending', 'processing'])
        ).all()
        
        results = []
        for candidate in pending:
            # Calculate progress
            progress = 0
            if candidate.interview_ai_analysis_status == 'processing':
                progress = 50
            
            results.append({
                'candidate_id': candidate.id,
                'name': candidate.name,
                'status': candidate.interview_ai_analysis_status,
                'progress': progress,
                'completed_at': candidate.interview_completed_at.isoformat()
            })
        
        return jsonify({
            'success': True,
            'pending_analyses': results
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/complete/<token>', methods=['POST', 'OPTIONS'])
def complete_interview(token):
    """Complete an interview and trigger AI analysis"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Update completion status
        candidate.interview_completed_at = datetime.now()
        candidate.interview_status = 'completed'
        
        # Calculate duration
        if candidate.interview_started_at:
            duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
            candidate.interview_duration = int(duration)
        
        # Set analysis status to pending
        candidate.interview_ai_analysis_status = 'pending'
        
        # Trigger AI analysis (you'll need to implement this)
        trigger_ai_analysis(candidate.id)
        
        session.commit()
        
        # Clear cache to ensure fresh data
        cache.delete_memoized(get_cached_candidates)
        
        return jsonify({
            "success": True,
            "message": "Interview completed successfully"
        }), 200
        
    finally:
        session.close()

def trigger_ai_analysis(candidate_id):
    """Trigger AI analysis for completed interview"""
    def run_analysis():
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return
            
            # Update status
            candidate.interview_ai_analysis_status = 'processing'
            session.commit()
            
            # Parse Q&A data
            questions = json.loads(candidate.interview_questions_asked or '[]')
            answers = json.loads(candidate.interview_answers_given or '[]')
            
            # Simulate AI analysis (replace with actual AI logic)
            # For now, generate random scores for testing
            import random
            
            candidate.interview_ai_score = random.randint(60, 95)
            candidate.interview_ai_technical_score = random.randint(60, 95)
            candidate.interview_ai_communication_score = random.randint(60, 95)
            candidate.interview_ai_problem_solving_score = random.randint(60, 95)
            candidate.interview_ai_cultural_fit_score = random.randint(60, 95)
            
            candidate.interview_ai_overall_feedback = f"""
Based on the interview analysis:
- Technical Skills: Strong understanding of core concepts
- Communication: Clear and articulate responses
- Problem Solving: Demonstrated analytical thinking
- Cultural Fit: Good alignment with company values

Total Questions Asked: {len(questions)}
Total Questions Answered: {len(answers)}
Completion Rate: {(len(answers)/len(questions)*100) if questions else 0:.1f}%
"""
            
            # Set final status
            if candidate.interview_ai_score >= 70:
                candidate.interview_final_status = 'Recommended'
            else:
                candidate.interview_final_status = 'Not Recommended'
            
            candidate.interview_ai_analysis_status = 'completed'
            session.commit()
            
            # Clear cache
            cache.delete_memoized(get_cached_candidates)
            
        except Exception as e:
            logger.error(f"AI analysis failed for candidate {candidate_id}: {e}")
            if candidate:
                candidate.interview_ai_analysis_status = 'failed'
                session.commit()
        finally:
            session.close()
    
    # Run in background thread
    executor.submit(run_analysis)

# Add this after your existing track_qa_enhanced function

@app.route('/api/interview/qa/get/<int:candidate_id>', methods=['GET'])
def get_interview_qa_data(candidate_id):
    """Get complete Q&A data for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Parse all Q&A data
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        qa_pairs = json.loads(getattr(candidate, 'interview_qa_pairs', '[]'))
        
        return jsonify({
            "success": True,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "qa_data": {
                "questions": questions,
                "answers": answers,
                "qa_pairs": qa_pairs,
                "total_questions": len(questions),
                "total_answers": len(answers),
                "completion_rate": f"{(len(answers) / len(questions) * 100) if questions else 0:.1f}%"
            },
            "transcript": candidate.interview_transcript,
            "analysis": {
                "status": candidate.interview_ai_analysis_status,
                "overall_score": candidate.interview_ai_score,
                "technical_score": candidate.interview_ai_technical_score,
                "communication_score": candidate.interview_ai_communication_score,
                "problem_solving_score": candidate.interview_ai_problem_solving_score,
                "cultural_fit_score": candidate.interview_ai_cultural_fit_score,
                "feedback": candidate.interview_ai_overall_feedback,
                "recommendation": candidate.interview_final_status
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting Q&A data: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/debug/candidate-fields/<int:candidate_id>', methods=['GET'])
def debug_candidate_fields(candidate_id):
    """Debug endpoint to see what fields a candidate has"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Get all attributes
        fields = {}
        for key in dir(candidate):
            if not key.startswith('_') and not callable(getattr(candidate, key)):
                try:
                    value = getattr(candidate, key)
                    fields[key] = str(type(value).__name__)
                except:
                    fields[key] = "error reading"
        
        return jsonify({
            "candidate_id": candidate_id,
            "available_fields": fields,
            "has_knowledge_base_id": hasattr(candidate, 'knowledge_base_id'),
            "has_interview_kb_id": hasattr(candidate, 'interview_kb_id')
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/recording/<int:candidate_id>', methods=['GET'])
def get_interview_recording_info(candidate_id):
    """Get interview recording information"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        recording_info = {
            "recording_file": candidate.interview_recording_file,
            "recording_duration": candidate.interview_recording_duration,
            "recording_size": candidate.interview_recording_size,
            "recording_format": candidate.interview_recording_format,
            "recording_quality": candidate.interview_recording_quality,
            "recording_status": candidate.interview_recording_status,
            "session_id": candidate.interview_session_id,
            "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
            "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None
        }
        
        return jsonify({
            "success": True,
            "candidate_id": candidate_id,
            "recording_info": recording_info
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting recording info: {e}")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        session.close()

@app.route('/api/routes', methods=['GET'])
def list_routes():
    """Debug endpoint to list all available routes"""
    routes = []
    for rule in app.url_map.iter_rules():
        if rule.endpoint != 'static':
            routes.append({
                'endpoint': rule.endpoint,
                'methods': list(rule.methods),
                'rule': str(rule)
            })
    return jsonify({
        "total_routes": len(routes),
        "routes": sorted(routes, key=lambda x: x['rule'])
    }), 200

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({
        "error": "Endpoint not found",
        "available_endpoints": [
            "/",
            "/api/jobs",
            "/api/candidates",
            "/api/run_full_pipeline",
            "/api/pipeline_status",
            "/api/recruitment-stats",
            "/secure-interview/<token>",
            "/health",
            "/api/routes"
        ]
    }), 404

def run_bulk_scraping_with_monitoring():
    """Wrapper to run bulk scraping with monitoring"""
    start_time = time.time()
    
    try:
        logger.info("Starting bulk scraping for all pending assessments")
        
        # Import and run the bulk scraping function
        try:
            from testlify_results_scraper import scrape_all_pending_assessments
        except ImportError as e:
            logger.error(f"Failed to import scraper: {e}")
            notify_admin(
                "Scraper Import Error",
                f"Could not import results scraper: {str(e)}. Please ensure testlify_results_scraper.py is available."
            )
            return
        
        # Run the async scraping function
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            results_summary = loop.run_until_complete(scrape_all_pending_assessments())
        finally:
            loop.close()
        
        duration = time.time() - start_time
        total_candidates = sum(results_summary.values()) if isinstance(results_summary, dict) else 0
        
        logger.info(f"Bulk scraping completed in {duration:.2f} seconds. Processed {len(results_summary)} assessments, {total_candidates} candidates.")
        
        # Send success notification
        if isinstance(results_summary, dict):
            summary_text = "\n".join([f"- {assessment}: {count} candidates" for assessment, count in results_summary.items()])
        else:
            summary_text = f"Processed {total_candidates} total candidates"
            
        notify_admin(
            "Bulk Assessment Results Scraping Completed",
            f"Assessments processed: {len(results_summary) if isinstance(results_summary, dict) else 'Unknown'}\nTotal candidates: {total_candidates}\nDuration: {duration:.2f} seconds\n\nBreakdown:\n{summary_text}"
        )
        
    except Exception as e:
        duration = time.time() - start_time
        error_msg = f"Bulk scraping failed after {duration:.2f} seconds"
        logger.error(error_msg, exc_info=True)
        
        # Send failure notification
        notify_admin(
            "Bulk Assessment Results Scraping Failed",
            error_msg,
            error_details=traceback.format_exc()
        )

@app.route('/health', methods=['GET'])
def health_check():
    """Enhanced health check endpoint with system status"""
    health_status = {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "2.1.0",
        "checks": {},
        "performance": {
            "avg_response_time": f"{request_metrics['avg_response_time']:.3f}s",
            "total_requests": request_metrics['total_requests'],
            "slow_requests": request_metrics['slow_requests']
        }
    }
    
    # Check database
    try:
        session = SessionLocal()
        session.execute("SELECT 1")
        session.close()
        health_status["checks"]["database"] = "healthy"
    except Exception as e:
        health_status["checks"]["database"] = f"unhealthy: {str(e)}"
        health_status["status"] = "degraded"
    
    # Check cache
    try:
        cache.set('health_check', 'ok', timeout=1)
        if cache.get('health_check') == 'ok':
            health_status["checks"]["cache"] = "healthy"
        else:
            health_status["checks"]["cache"] = "unhealthy"
            health_status["status"] = "degraded"
    except Exception as e:
        health_status["checks"]["cache"] = f"unhealthy: {str(e)}"
        health_status["status"] = "degraded"
    
    # Check thread pool
    try:
        if hasattr(executor, '_threads'):
            active_threads = len([t for t in executor._threads if t.is_alive()])
            health_status["checks"]["thread_pool"] = f"healthy ({active_threads} active threads)"
        else:
            health_status["checks"]["thread_pool"] = "healthy"
    except Exception as e:
        health_status["checks"]["thread_pool"] = f"degraded: {str(e)}"
    
    return jsonify(health_status), 200 if health_status["status"] == "healthy" else 503

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Endpoint not found","available_endpoints":["/","/api/jobs","/api/candidates","/api/secure_interview/<token>","/heath"]}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {error}")
    return jsonify({"error": "Internal server error"}), 500

@app.errorhandler(429)
def rate_limit_exceeded(error):
    return jsonify({"error": "Rate limit exceeded. Please try again later."}), 429

# Request logging
@app.before_request
def log_request_info():
    """Log incoming requests"""
    logger.info(f"🌐 {request.method} {request.path} from {request.remote_addr}")

def test_routes():
    """Test if routes are properly registered"""
    print("📋 Registered Routes:")
    for rule in app.url_map.iter_rules():
        print(f"  {rule.methods} {rule.rule} -> {rule.endpoint}")




# Cleanup on shutdown
import atexit

def cleanup():
    """Cleanup resources on shutdown"""
    logger.info("Shutting down TalentFlow AI Backend...")
    executor.shutdown(wait=True)
    
atexit.register(cleanup)

if __name__ == "__main__":
    print("Starting TalentFlow AI Backend (Optimized Version)...")
    print("Server running at http://127.0.0.1:5000")
    print("Logging to: logs/talentflow.log")
    print("Performance optimizations enabled")
    print("Caching enabled")
    print("Pipeline status tracking enabled")

    print("Starting Interview Automation System...")
    start_interview_automation()
    print("Interview automation running (checking every 30 minutes)")


    with app.app_context():
        print("\n Registered Routes:")
        for rule in app.url_map.iter_rules():
            print(f"  {list(rule.methods)} {rule.rule} -> {rule.endpoint}")

    try:
        try: 
            print("\n🤖 Starting Interview Automation System...")
            start_interview_automation()
            print("✅ Interview automation running (checking every 30 minutes)")
        except Exception as e:
            print(f"⚠️  Interview automation not available: {e}")
            


        app.run(
            host='0.0.0.0',
            port=5000,
            debug=os.getenv('FLASK_ENV') == 'development',
            use_reloader=False,
            threaded=True
        )
    finally:
        print("Shutting down interview automation...")
        stop_interview_automation()